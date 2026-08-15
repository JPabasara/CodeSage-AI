# -*- coding: utf-8 -*-
"""
⚠️ HISTORICAL. This script built the v0.x SKELETON from the blank course template.
It is superseded by the patch approach in _patch_srs_v1_1.py / _patch_sad_v1_1.py,
which edit the previous version instead of rebuilding from the template - because
rebuilding throws away every figure and hand-written table added since.

Kept because it is the record of how the document structure was first laid out,
and because it is the only thing that knows how to inherit the template's styles,
headers, footers and numbering from scratch. See docs/tools/README.md.

Build the Code Sage AI SRS .docx ON TOP OF the course template so that every
style, header, footer, margin and numbering definition is literally the
template's own (styles.xml / numbering.xml / header*.xml / footer*.xml / sectPr
are all inherited untouched).

Body content is rebuilt; blue bracketed text = guidance/TO FILL (delete before
submission), per the template's own convention.
"""
import re, shutil, zipfile, os
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ───────────────────────────────────────────────────────────────────
# Repo-relative, so the script runs on any machine and from any directory.
# The course templates stayed in docs/Templates/ when the scripts moved here.
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
TPL = os.path.join(ROOT, "docs", "Templates", "3 = Template for Software Requirements Specification.docx")
OUT = os.path.join(ROOT, "docs", "Deliverables", "_generated", "Software_Requirements_Specification_skeleton.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

# Writing into _generated/ rather than over a deliverable is deliberate: this
# script produces a SKELETON from the blank course template, not a finished
# document. Pointing it at the real filename would make one careless run replace
# a document the team had spent days editing.
if os.path.exists(OUT):
    raise SystemExit(
        f"{OUT} already exists - delete it first if you really want to rebuild "
        "the skeleton from the blank course template."
    )



TBLPR = (
    '<w:tblPr %s>'
    '<w:tblW w:w="0" w:type="auto"/>'
    '<w:tblBorders>'
    '<w:top w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:left w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:right w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
    '<w:tblLayout w:type="fixed"/>'
    '<w:tblLook w:val="0000" w:firstRow="0" w:lastRow="0" w:firstColumn="0" '
    'w:lastColumn="0" w:noHBand="0" w:noVBand="0"/>'
    '</w:tblPr>' % nsdecls('w')
)

shutil.copyfile(TPL, OUT)
doc = docx.Document(OUT)
body = doc.element.body

# ───────────────────────── 1. placeholder replacement (body + hdr/ftr) ────────
REPL = {
    "<Project Name>": "Code Sage AI",
    "<Company Name>": "Group 16",
    "<Subsystem or Feature>": "the AI-Powered Technical-Debt Analytics Dashboard",
    "<1.0>": "1.0",
    "<dd/mmm/yy>": "29/Jul/2026",
    "<document identifier>": "CS3203-G16-SRS-v1.0",
}


def fix_runs(container):
    for p in container.paragraphs:
        for r in p.runs:
            for k, v in REPL.items():
                if k in r.text:
                    r.text = r.text.replace(k, v)
    for t in container.tables:
        for row in t.rows:
            for c in row.cells:
                fix_runs(c)


fix_runs(doc)
# NOTE: headers/footers are deliberately NOT touched through python-docx here.
# Reading section.first_page_header / even_page_footer CREATES an empty definition
# and breaks inheritance from the template. They are patched at the XML level in
# the post-processing step at the bottom of this script instead.

doc.core_properties.title = "Software Requirements Specification"
doc.core_properties.subject = "Code Sage AI"
doc.core_properties.author = "Group 16"
doc.core_properties.category = "CS3203 Software Engineering Project"

# ───────────────────────── 2. index the original body blocks ─────────────────
kids = list(body.iterchildren())

# title-page guidance notes (blocks 8, 9) -> remove
for i in (8, 9):
    body.remove(kids[i])

# group identification on the title page, after "Version 1.0"
anchor = kids[6]
for line, sz, bold in [
    ("Group 16   |   Project ID 7   |   CS3203 - Software Engineering Project", 12, True),
    ("Mentor: Mr. Anju Chamantha", 11, False),
    ("Date: 29 July 2026", 11, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(sz)
    r.bold = bold
    anchor.addnext(p._p)
    anchor = p._p

# ───────────────────────── 3. Revision History table ─────────────────────────
rev = doc.tables[0]
rows = [
    ("22/Jul/2026", "0.1", "Initial draft: v1.0 functional and non-functional requirements, interfaces, database and constraints.", "Group 16"),
    ("29/Jul/2026", "0.2", "Restructured onto the course SRS template; Sections 3.2, 3.5 and 3.9 expanded with measurable requirements.", "Group 16"),
    ("30/Jul/2026", "0.3", "CR-001: severity register (FR-8.1), source reduced to rule|satd (FR-8.2), SATD marker severity (FR-9.2), "
                           "risk as a bounded multiplier (FR-10/FR-11), profile = 5 category weights + trust slider (FR-20), "
                           "in-place finding detail (FR-17/FR-18), three-mechanism visibility floor (FR-24), Appendix C completed.", "Group 16"),
    ("31/Jul/2026", "0.4", "CR-001 (D-CR8 to D-CR12): a profile change is not a scan and writes no snapshot (FR-20); snapshots "
                           "store facts while all scores are derived on read (FR-21, DB-8); the trend chart uses one lens and is "
                           "labelled with the active profile (FR-14); debt-category taxonomy confirmed against the SATD dataset - "
                           "defect added as a sixth category, D5 closed (FR-9.3).", "Group 16"),
]
for ri, data in enumerate(rows, start=1):
    for ci, val in enumerate(data):
        cell = rev.rows[ri].cells[ci]
        para = cell.paragraphs[0]
        para.style = doc.styles["Tabletext"]
        if para.runs:
            para.runs[0].text = val
            for extra in para.runs[1:]:
                extra.text = ""
        elif val:
            para.add_run(val)

# ───────────────────────── 4. TOC: fresh auto-updating field ─────────────────
kids = list(body.iterchildren())
toc_start = next(i for i, e in enumerate(kids)
                 if e.tag == qn('w:p') and 'TOC \\o' in e.xml)
# The whole cached TOC (entries + the paragraph holding the outer field end) sits
# between the "Table of Contents" title and the first Heading 1. Individual entries
# each carry their own inner PAGEREF field end, so scan for the heading instead.
intro = next(i for i, e in enumerate(kids)
             if e.tag == qn('w:p') and 'Heading1' in e.xml and 'Introduction' in e.xml)
after = kids[toc_start - 1]
for e in kids[toc_start:intro]:
    body.remove(e)

toc_xml = (
    '<w:p %s><w:pPr><w:pStyle w:val="TOC1"/></w:pPr>'
    '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
    '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    '<w:r><w:t xml:space="preserve">Right-click here and choose "Update Field" '
    '(or press Ctrl+A then F9) to generate the Table of Contents.</w:t></w:r>'
    '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>' % nsdecls('w')
)
after.addnext(parse_xml(toc_xml))

# ───────────────────────── 5. wipe the old body content ──────────────────────
kids = list(body.iterchildren())
start = next(i for i, e in enumerate(kids)
             if e.tag == qn('w:p') and 'Introduction' in e.xml and 'Heading1' in e.xml)
for e in kids[start:]:
    if e.tag != qn('w:sectPr'):
        body.remove(e)

# ───────────────────────── 6. content helpers ────────────────────────────────
def h1(t):
    doc.add_paragraph(t, style="Heading 1")


def h2(t):
    doc.add_paragraph(t, style="Heading 2")


def h3(t):
    doc.add_paragraph(t, style="Heading 3")


def info(t):
    doc.add_paragraph(t, style="InfoBlue")


def body_text(t, bold_lead=None):
    p = doc.add_paragraph(style="Body Text")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(t)
    return p


def bullet(t, lead=None):
    p = doc.add_paragraph(style="Bullet1")
    p.add_run("\u2022\t")
    if lead:
        r = p.add_run(lead)
        r.bold = True
    p.add_run(t)
    return p


def numitem(n, t):
    p = doc.add_paragraph(style="Bullet1")
    p.add_run("%s\t" % n)
    p.add_run(t)
    return p


def caption(t):
    p = doc.add_paragraph(style="Body Text")
    r = p.add_run(t)
    r.bold = True
    r.italic = True
    return p


def table(headers, data, widths, cap=None):
    if cap:
        caption(cap)
    t = doc.add_table(rows=1 + len(data), cols=len(headers))
    tbl = t._tbl
    tbl.replace(tbl.tblPr, parse_xml(TBLPR))
    grid = tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(w))

    def fill(cell, text, bold=False, center=False):
        cell._tc.remove(cell._tc.find(qn('w:p')))
        for i, line in enumerate(str(text).split("\n")):
            p = cell.add_paragraph(style="Tabletext")
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # **bold** segments
            for j, seg in enumerate(re.split(r"\*\*(.+?)\*\*", line)):
                if seg:
                    r = p.add_run(seg)
                    r.bold = bold or (j % 2 == 1)

    for ci, htxt in enumerate(headers):
        c = t.rows[0].cells[ci]
        c.width = docx.shared.Twips(widths[ci])
        fill(c, htxt, bold=True, center=True)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml('<w:tblHeader %s/>' % nsdecls('w')))
    for ri, rowdata in enumerate(data, start=1):
        for ci, val in enumerate(rowdata):
            c = t.rows[ri].cells[ci]
            c.width = docx.shared.Twips(widths[ci])
            fill(c, val)
    doc.add_paragraph(style="Body Text")
    return t


W2 = [2300, 7060]
W3 = [1500, 3900, 3960]
W4 = [900, 3300, 3060, 2100]

# ═════════════════════════ 1. INTRODUCTION ═══════════════════════════════════
h1("Introduction")
info("[Guidance: the Introduction gives an overview of the whole SRS - purpose, "
     "scope, glossary, references and document organisation. Blue bracketed text "
     "throughout this document is guidance or a TO FILL marker and must be deleted "
     "before submission. Search for \"TO FILL\" to find every outstanding item.]")

h2("Purpose")
body_text("This document specifies the complete software requirements for Code Sage AI, "
          "an AI-powered technical-debt analytics dashboard delivered as a multi-tenant web "
          "platform for small-scale agile software teams. It defines the system's external "
          "behaviour - functional capabilities, non-functional qualities, interfaces, data and "
          "design constraints - in enough detail for designers to design the system and for "
          "testers to verify it.")
body_text("The intended audience is the development team (frontend, backend and machine-learning "
          "members), the project mentor and evaluators, and any future maintainer. The requirements "
          "below are the agreement against which the v1.0 release is built and tested; "
          "later-release requirements are included so that the product trajectory is visible and "
          "the v1.0 design stays forward-compatible.")

h2("Scope")
body_text("Code Sage AI connects a team's GitHub repository, analyses its source code, code comments "
          "and commit history using static analysis plus machine learning, and presents a prioritized, "
          "low-noise view of the codebase's technical debt: an overall code-health score and grade, a "
          "hotspot file-tree heat map, a \"Refactor-First\" list of the issues most worth fixing, and "
          "health trends over time. Unlike enterprise tools that overwhelm small teams with hundreds "
          "of findings and heavy configuration, Code Sage AI surfaces only the most critical, "
          "actionable items and explains each in one plain-English line.")
body_text("This SRS specifies the v1.0 release. Every requirement carries a release tag; requirements "
          "tagged [v1.1] or [v2] are specified now to protect the design but are not built for v1.0.")
table(
    ["Release tag", "Meaning"],
    [("[v1.0]", "Built, tested and demonstrated in this release."),
     ("[v1.1]", "Specified now, built in the next release. The v1.0 design must not block it."),
     ("[v2]", "Specified now, built in the Team-tier release. Architectural seam only in v1.0."),
     ("[supporting]", "Not a runtime feature; produced as project evidence (e.g. model evaluation).")],
    W2, "Table 1-1. Release tags used throughout this document.")
table(
    ["Category", "Items"],
    [("In scope (v1.0)",
      "GitHub sign-in; connecting one public repository by URL (1 repo = 1 project); a Projects list "
      "with selection; per-branch on-demand scanning with progress and cancel; the analysis pipeline "
      "(Lizard metrics, PyDriller history, rule engine, SATD classifier, risk model); weighted scoring "
      "with preset profiles **and custom weight sliders**, and a critical-security visibility floor; the "
      "dashboard outputs (health card with category pie, health trend chart, Refactor-First list with "
      "debt-type filter, in-place finding detail, hotspot file-tree heat map); scan history; persistence "
      "of every scan as an immutable snapshot; a single-member workspace built on a multi-tenant foundation."),
     ("Deferred ([v1.1])",
      "Private repositories via GitHub App; finding actions (accept / resolve / false-positive); code "
      "snippet on demand; standalone category-breakdown view; graduated rule severity; separate trust "
      "dials for the two machine-learning models; accessibility hardening pass."),
     ("Deferred ([v2])",
      "Multi-repository workspaces; Team tier and RBAC; silent checks (auto-scan on push or pull "
      "request); cross-repository dependency analytics; GitLab support; billing."),
     ("Not detection inputs in v1.0",
      "Commit-message text as a SATD input (history is consumed only as four numeric process metrics); "
      "pull requests and issues in any form; previously stored snapshots as a model or scoring input "
      "(see FR-7.1)."),
     ("Out of scope (whole product)",
      "Fully automatic code refactoring; on-premise or proprietary version-control systems beyond "
      "GitHub and GitLab; integration with project-management tools such as Jira or Trello; full "
      "vulnerability scanning (SAST/DAST, CVE and dependency auditing, penetration testing).")],
    W2, "Table 1-2. Scope of the v1.0 release.")

h2("Definitions, Acronyms, and Abbreviations")
table(
    ["Term", "Meaning"],
    [("Technical debt (TD)", "The implied future cost of shortcuts or poor design decisions taken in code."),
     ("SATD", "Self-Admitted Technical Debt - debt a developer admits in natural language (for example "
              "\"// TODO: temporary hack\"). The literature recognises four sources (comments, commit "
              "messages, issues, pull requests); v1.0 detects SATD in source-code comments only (FR-9.1)."),
     ("Finding", "The atomic unit of output: one detected issue at a file:line:symbol, with a source, a "
                 "category, a severity and a one-line reason."),
     ("source", "Which detector produced a finding: **rule | satd** - the only two producers. Security patterns run "
                "inside the rule engine, so a security finding is a rule finding whose category is security; the risk "
                "model produces no findings at all (FR-8.2)."),
     ("category", "What type of debt a finding is: code-design | requirement | documentation | test | security."),
     ("severity", "How bad a finding is: critical | high | medium | low. Assigned at detection - from the rule register "
                  "for rule findings, from the comment-marker table for SATD findings (FR-9.2) - and never by scoring, "
                  "by the user interface, by a machine-learning model, or by a user (FR-8.1). Stored on the finding and "
                  "read by exactly two consumers: scoring (base points) and the Refactor-First badge."),
     ("Trust slider (s)", "A single profile control, s in [0,1], default 0.5, expressing how much a team trusts the "
                          "deterministic rules versus the machine-learning detectors. Yields rule_trust = 0.5 + s and "
                          "ml_trust = 1.5 - s; the security category is excluded and always uses 1.0 (FR-11, FR-24)."),
     ("severity", "How bad a finding is: critical | high | medium | low. Assigned at detection by the detector "
                  "that fired - never by scoring, never by the client, never by a machine-learning model "
                  "(FR-8.1) - then read by exactly two consumers: scoring, which maps it to base points "
                  "(FR-11), and the Refactor-First badge (FR-15)."),
     ("Rule engine", "Deterministic thresholds over static metrics plus pattern-based security rules."),
     ("Risk model (ML-2)", "Supervised classifier estimating a file's bug-proneness (0-1)."),
     ("SATD classifier (ML-1)", "Supervised NLP model classifying a source-code comment as debt and, if so, its category."),
     ("CCN", "Cyclomatic Complexity Number (per function)."),
     ("NLOC", "Non-comment Lines Of Code."),
     ("Churn", "Recent change volume of a file, over a 90-day window anchored to the scanned commit's date (FR-11)."),
     ("Process metrics", "The four history-derived numeric features per file: churn, author count, file age, "
                         "recency. Mined by PyDriller; consumed by ML-2 and by scoring (churn only)."),
     ("Snapshot", "The immutable stored result of one scan, keyed by repository, branch, commit SHA and timestamp."),
     ("Snapshot-scoped", "Derived solely from the working tree at the scanned commit SHA."),
     ("Health score / grade", "0-100 score and A-E grade summarising a repository's or subtree's debt."),
     ("Scoring profile", "The user-owned settings that turn findings into scores: six per-category weights plus the "
                         "trust slider s (FR-20). It never contains a severity."),
     ("Visibility floor", "The rule that critical security findings are never suppressed or down-weighted below visibility."),
     ("Workspace / tenant", "The top-level data owner in the multi-tenant model; a project belongs to a workspace, not a person."),
     ("RBAC", "Role-Based Access Control (org-admin / manager / developer / viewer)."),
     ("RLS", "PostgreSQL Row-Level Security, used to isolate tenants."),
     ("SZZ", "Algorithm that links bug-fixing commits back to the changes that induced them (used to label defective files)."),
     ("GHPR", "GitHub Pull Request defect dataset, used for risk-model training and evaluation. Offline "
              "training data only - the system never reads pull requests at scan time."),
     ("MSW", "Mock Service Worker - intercepts network calls to serve mock API data during frontend development and testing."),
     ("FR / NFR", "Functional / Non-Functional Requirement."),
     ("SAD / SDD", "Software Architecture Document / Software Design Document.")],
    W2, "Table 1-3. Definitions, acronyms and abbreviations.")

h2("References")
info("[Guidance: list every referenced document in IEEE style - title, version, date, source. "
     "TO FILL: import the full numbered reference list from the Project Proposal and Feasibility "
     "Report and align the citation numbers across all deliverables.]")
for n, t in [
    ("[1]", "Code Sage AI - Project Proposal, Group 16, 5 July 2026."),
    ("[2]", "Code Sage AI - Feasibility Report, Group 16, 12 July 2026."),
    ("[3]", "Code Sage AI - Backend Analysis Engine: Detection, Scoring and Output Generation, v1.0, 21 July 2026."),
    ("[4]", "Code Sage AI - Release Roadmap and Frontend Prototype Plan, Group 16, July 2026."),
    ("[5]", "Code Sage AI - Software Architecture Document, Group 16, 22 July 2026."),
    ("[6]", "T. Besker, A. Martini and J. Bosch, \"Technical Debt Cripples Software Developer Productivity,\" 2018."),
    ("[7]", "Y. Li, M. Soliman and P. Avgeriou, \"Automatic identification of self-admitted technical debt from "
            "four different sources,\" Empirical Software Engineering, vol. 28, no. 65, 2023."),
    ("[8]", "J. Xu, F. Wang and J. Ai, \"Defect Prediction With Semantics and Context Features,\" "
            "IEEE Transactions on Reliability, 2021."),
    ("[9]", "IEEE Std 830-1998, Recommended Practice for Software Requirements Specifications."),
]:
    numitem(n, t)
doc.add_paragraph(style="Body Text")

h2("Overview")
body_text("Section 2 gives the overall description: product perspective, product functions, user "
          "characteristics, constraints and assumptions. Section 3 gives the specific requirements - "
          "3.1 functionality (grouped by feature and described in detail, without use-case diagrams), "
          "followed by the non-functional requirements in 3.2 usability, 3.3 reliability, 3.4 performance "
          "and security and 3.5 supportability, then 3.6 design constraints and 3.7 to 3.12 covering "
          "documentation, purchased components, interfaces, database, legal notices and applicable "
          "standards. Section 4 holds supporting information and appendices.")

# ═════════════════════════ 2. OVERALL DESCRIPTION ════════════════════════════
h1("Overall Description")
info("[Guidance: general factors and background - not specific requirements. Cover product "
     "perspective, product functions, user characteristics, constraints, assumptions and "
     "dependencies, and requirements subsets.]")

body_text("Code Sage AI is a new, self-contained web product; it is not a component of a larger system. "
          "It integrates externally with the GitHub API and with Git itself to read repositories, and is "
          "built from four internal parts: a Next.js frontend (the dashboard), a FastAPI backend with "
          "Celery and Redis asynchronous workers (repository integration, orchestration, scoring), a "
          "Python/scikit-learn machine-learning service (SATD and risk models over Lizard metrics and "
          "PyDriller history), and a PostgreSQL database (tenant data and immutable scan snapshots). A "
          "load-bearing architectural invariant is that the dashboard computes nothing: every number "
          "shown traces to a stored row, and all detection and scoring happen server-side.",
          bold_lead="Product perspective. ")
body_text("Connect a repository, scan a branch on demand (asynchronously), extract signals, detect "
          "findings (rules, SATD and risk), score and prioritize them with a profile, persist an "
          "immutable snapshot, and present the health score, grade and delta, a category pie, a health "
          "trend, a Refactor-First list with one-line reasons and a debt-type filter, a finding-detail "
          "panel, and a hotspot file-tree heat map; browse scan history; select a scoring profile.",
          bold_lead="Product functions. ")

table(
    ["User class", "Description", "Primary needs"],
    [("Developer (primary)", "Connects repositories, runs scans, reads and triages findings. Comfortable "
                             "with code.", "Low noise, actionable items, evidence for each finding."),
     ("Tech lead / Manager", "Configures the scoring profile for the team and watches health trends.",
      "Prioritization they can steer; trend over time to plan refactoring within sprints."),
     ("Viewer / stakeholder", "Product owner or non-technical business owner reading codebase condition.",
      "Jargon-light presentation; score, grade and trend understandable without reading code."),
     ("Org admin [v2]", "Manages the workspace, connected repositories, members, roles and billing.",
      "Access control and least-privilege repository access.")],
    W3, "Table 2-1. User characteristics.")

body_text("A three-person team on a fixed academic schedule (core features first); free tools, "
          "datasets and hosting only, with no budget; GitHub API rate limits mitigated with ETag caching "
          "and App tokens; machine-learning validated on Python and JavaScript/TypeScript only; "
          "repositories read through least-privilege, user-selected access.",
          bold_lead="Constraints. ")
body_text("The system assumes a reachable Git host, that the target repository is in a supported "
          "language for full detection quality, and that the public training datasets remain available "
          "under their current licences.",
          bold_lead="Assumptions and dependencies. ")
body_text("Each functional requirement carries [v1.0], [v1.1] or [v2]. v1.0 is a complete vertical slice "
          "- one user, one public repository, manual scan, full dashboard - that exercises every "
          "architectural layer.", bold_lead="Requirements subsets. ")

# ═════════════════════════ 3. SPECIFIC REQUIREMENTS ══════════════════════════
h1("Specific Requirements")
info("[Guidance: this section contains all software requirements at a level of detail sufficient for "
     "designers to design a system satisfying them and for testers to verify that they are satisfied.]")

# ---------- 3.1 Functionality ----------
h2("Functionality")
info("[Guidance: do not use use-case diagrams. Describe each function of the system / user activity in "
     "detail - trigger, behaviour, inputs and outputs. The table below is the requirement register; "
     "3.1.1 to 3.1.3 are worked examples of the level of detail expected. TO FILL: expand the remaining "
     "functional requirements into their own numbered sub-sections using the same pattern.]")
body_text("Notation: each requirement is written as FR-n [release] Name. \"The system\" means Code Sage AI.")

table(
    ["ID", "Release", "Requirement", "Summary"],
    [("FR-1", "v1.0", "Authentication and session",
      "Sign in with GitHub; authenticated session; land on Projects; sign out from the Account menu. "
      "Session state is carried per request (token), never held in server memory."),
     ("FR-2", "v1.0 seam / v2 full", "Workspace and tenant isolation",
      "Every project, scan and finding belongs to a workspace (tenant), not to a user. v1.0 has one "
      "member per workspace; workspace_id and PostgreSQL Row-Level Security are in place from the first "
      "migration so Team/RBAC attaches with no schema rewrite."),
     ("FR-3", "v1.0", "Connect a repository (public URL)",
      "Add a project by pasting a public repository URL (1 repository = 1 project). The system validates "
      "the URL and records name, owner, visibility and default branch. FR-3b [v1.1]: private repositories "
      "through a GitHub App installation with user-selected repositories."),
     ("FR-4", "v1.0", "Projects list and selection",
      "Display connected projects vertically with name, owner, visibility and a health hint; Select sets "
      "the active project, which scopes the Dashboard and Scan-History views."),
     ("FR-5", "v1.0", "Branch selection",
      "Branch dropdown in the dashboard top navigation, default branch pre-selected. Analysis is "
      "per-branch: each branch has its own snapshots and trend."),
     ("FR-6", "v1.0", "Repository scan (manual, asynchronous, cancellable)",
      "Scan control runs the pipeline on an asynchronous worker: idle -> queued -> running NN% (with Stop) "
      "-> done or error. Cancel leaves the previous snapshot intact. The system may skip re-scanning when "
      "the branch head SHA equals the last scanned SHA. FR-6b [v2]: silent checks on push or pull request."),
     ("FR-7", "v1.0", "Signal extraction",
      "From a local clone at the selected branch's commit SHA: static metrics with Lizard; process metrics "
      "with PyDriller (churn, author count, file age, recency); and source-code comments from the working "
      "tree at that SHA for the SATD classifier."),
     ("FR-7.1", "v1.0", "Extraction boundary (normative)",
      "Git history enters the pipeline as aggregated numbers, never as text. Commit-message text, pull "
      "requests, issues and previously stored snapshots are not detection inputs. Consequence: a scan is a "
      "pure function of the repository at that SHA, a fixed model version and the active profile."),
     ("FR-8", "v1.0", "Detection - rule engine",
      "Deterministic thresholds (CCN > 15, method > 80 NLOC, nesting > 4, duplicated block, file > 800 NLOC) "
      "plus security patterns (hardcoded secret by regex and entropy, SQL string concatenation, dangerous "
      "eval/exec). Each finding records file, line, symbol, category, severity, measured value, threshold and "
      "rule id."),
     ("FR-8.1", "v1.0", "Severity and category assigned at detection (normative)",
      "Every rule definition carries a fixed category and a fixed severity alongside its message template; both "
      "are written onto the finding at detection and are never recomputed by scoring, by a profile change or by "
      "the client. The SATD classifier predicts category only - its severity comes from the marker table in "
      "FR-9.2; the risk model assigns neither. **No machine-learning model and no user assigns a severity**: "
      "severity answers \"how bad is this kind of problem\" (the same answer for every team) while the profile's "
      "category weight answers \"how much does this team care\" (different per team). The register is Appendix C."),
     ("FR-8.2", "v1.0", "source has exactly two values (normative)",
      "source shall be rule or satd - the only two producers of findings. Security patterns execute inside the "
      "rule engine, so a security finding is source = rule with category = security; a security source value "
      "would duplicate the category axis exactly and destroy its orthogonality. The risk model emits no findings, "
      "so no finding can carry an ml-risk source."),
     ("FR-9", "v1.0", "Detection - SATD classifier (ML-1)",
      "Classify each extracted source-code comment as debt or not and, if debt, assign a category from "
      "{code-design, requirement, defect, documentation, test} (FR-9.3), using a supervised NLP model trained on "
      "the Li et al. SATD dataset. Each finding is anchored to the comment's file:line and quotes the comment "
      "text. The model predicts category only; severity is assigned per FR-9.2."),
     ("FR-9.1", "v1.0", "SATD scope - the comments file only",
      "v1.0 trains, validates and evaluates on satd-dataset-code_comments.csv alone; the commit-message, issue "
      "and pull-request files are not used. Training and inference then share one distribution, so there is no "
      "train/serve skew. Commit-message SATD is excluded from detection because it has no file:line anchor, it "
      "accumulates forever and would invalidate the delta and trend, it carries no resolution signal, and - "
      "decisively - the four sources do not share a taxonomy (comments merge code and design into one label; the "
      "other three split them and add architecture_debt and build_debt)."),
     ("FR-9.3", "v1.0", "Debt-category taxonomy and label mapping (normative - closes D5)",
      "The taxonomy is fixed by the labels in the comments file (62,275 labelled comments): code-design <- "
      "code/design_debt (2,703), requirement <- requirement_debt (757), **defect <- defect_debt (472)**, test <- "
      "test_debt (85), documentation <- documentation_debt (54); plus security, which is not in the dataset and "
      "is emitted only by the rule engine. non_debt (58,204) is the negative class of the debt/not-debt decision "
      "and is never a category. Product values are normalised forms related to the dataset labels by a "
      "documented 1:1 mapping applied once in the ML service's post-processing - a deterministic rename cannot "
      "affect trainability. Only 6.54% of comments are debt, and documentation (0.09%) and test (0.14%) are "
      "severely under-represented, so per-class metrics are mandatory under FR-25."),
     ("FR-9.2", "v1.0", "SATD severity - the comment-marker table (normative)",
      "After the classifier determines that a comment is debt and assigns its category, severity is assigned by "
      "matching the comment text: FIXME/BUG/XXX/BROKEN -> high; TODO/HACK/TEMP/WORKAROUND -> medium; "
      "NOTE/REVIEW/NIT -> low; no marker -> medium. Patterns are evaluated high to low, highest match wins, and "
      "match anywhere in the comment. A supervised model can predict only what its training data labels, and the "
      "dataset labels categories, not severities - so severity cannot be learned and must be deterministic. The "
      "patterns and templates are in Appendix C.2."),
     ("FR-10", "v1.0", "Detection - risk model (ML-2)",
      "Per-file bug-proneness score (0-1) from a supervised classifier over Lizard product metrics plus the "
      "four PyDriller process metrics, trained on labelled defect data. It assigns neither a category nor a "
      "severity and produces no findings. It has exactly two effects: it boosts the priority of the findings in "
      "that file through the bounded risk_factor multiplier (FR-11), and it appears as a per-file risk badge. A "
      "risky file with no findings contributes no debt - every point of debt traces to a finding the user can open."),
     ("FR-11", "v1.0", "Scoring and prioritization",
      "A pure function over stored findings using the active profile. finding_priority = base_points(severity) x "
      "category_weight[category] x source_trust(finding) x churn_factor(file) x risk_factor(file); "
      "file_debt = sum of open finding priorities; repo_health = 100 x (1 - min(1, sum file_debt / (k x KLOC))); "
      "grade A >= 85, B >= 70, C >= 55, D >= 40, E < 40. churn_factor = 1 + min(commits_90d, 20)/20 (1.0-2.0); "
      "risk_factor = 1 + ml_trust x risk_score (1.0-2.5); rule_trust = 0.5 + s, ml_trust = 1.5 - s, and "
      "source_trust is 1.0 for the security category, rule_trust for rule findings and ml_trust for SATD findings. "
      "Base points are a lookup over the stored severity, not a judgement made at scoring time. The 90-day churn "
      "window is anchored to the scanned commit's committer date; wall-clock time is never used. The maximum "
      "combined boost (churn 2.0 x risk 2.5 = 5x) is less than the 8x Low-to-Critical spread, so within a category "
      "no model signal can raise a low finding above a critical one."),
     ("FR-12", "v1.0", "Overall Health card",
      "Health score (0-100), grade (A-E), delta versus the previous snapshot and a count of critical/high "
      "issues, for the selected branch."),
     ("FR-13", "v1.0", "Category breakdown pie",
      "Pie chart of technical debt by category, computed from stored findings. FR-13b [v1.1]: a standalone, "
      "filterable category-breakdown view."),
     ("FR-14", "v1.0", "Health trend chart",
      "Health per scan/commit over time for the selected branch, derived from the stored snapshots. **One lens per "
      "line:** every point is computed under the currently active profile, so selecting a different profile "
      "redraws the entire history under it and every point stays comparable; a line mixing profiles is prohibited. "
      "The chart is labelled with the active profile name. FR-14b [v2]: re-scope the chart to a hovered file or "
      "folder."),
     ("FR-15", "v1.0", "Refactor-First list and debt-type filter",
      "Prioritized list of rule and SATD findings sorted by priority; each row shows a category chip, a severity "
      "chip, file:line and the one-line reason. A row whose source is satd additionally carries a SATD chip; rule "
      "rows carry none, since rule is the default and a chip on every row carries no information. The source is "
      "shown as its own chip and never inside the severity chip, because severity is an ordinal scale on which a "
      "source value is not comparable and would leave base_points undefined. The badge renders the stored severity "
      "(FR-8.1); the client makes no severity judgement of its own. The user can filter by debt type."),
     ("FR-16", "v1.0", "One-line reason (deterministic templates)",
      "Every finding carries a one-line plain-English reason generated from string templates with the finding's "
      "own values interpolated - explainable, instant and incapable of hallucinating."),
     ("FR-17", "v1.0", "Finding detail - in-place detail mode",
      "Selecting a finding switches the dashboard into detail mode, rendered in place rather than as an overlay: "
      "the region holding the health card and trend chart is replaced by the finding's evidence, one-line reason "
      "and file:line:symbol (with room for the [v1.1] snippet); the file tree auto-expands and highlights that "
      "finding's file; and the Refactor-First list condenses to a strip so the user can move between findings. "
      "Closing restores the cards. The selected finding is reflected in the URL. FR-17b [v1.1]: code snippet on "
      "demand. FR-17c [v1.1]: accept-debt / resolve / false-positive actions - v1.0 is view-only."),
     ("FR-18", "v1.0", "Hotspot file-tree heat map",
      "Interactive file tree tinted red to green by health with a per-file risk badge, folders aggregating their "
      "children, expand/collapse and drill-in that re-aggregates from stored file scores without re-scanning. When "
      "the dashboard enters detail mode the tree automatically expands the ancestors of the selected finding's "
      "file, scrolls it into view and highlights it, remaining fully interactive throughout."),
     ("FR-19", "v1.0", "Scan history",
      "List past snapshots for the active project and branch (date, commit SHA, health score, grade, delta, "
      "finding count); selecting one loads that snapshot into the dashboard read-only."),
     ("FR-20", "v1.0", "Scoring profiles - presets and custom weights",
      "A profile consists of six per-category weights (security, code-design, defect, requirement, documentation, test; "
      "clamped 0.1-3.0) and one trust slider s (0-1, default 0.5). Three presets - Balanced (the default for every "
      "new workspace), Security-first and Delivery-speed - seed the sliders in one click, with a Reset to preset "
      "action. Any change re-scores the stored findings instantly with no re-scan. Weights are clamped because "
      "repo_health is calibrated against k and unbounded weights would drive every repository to grade E. The "
      "profile shall never contain a severity. **A profile change is not a scan:** it re-scores stored findings in "
      "place, does not run the pipeline, does not require the Scan control, and writes no snapshot - a snapshot is "
      "keyed by commit SHA and a profile is not a commit."),
     ("FR-21", "v1.0", "Snapshot persistence - facts stored, scores derived",
      "Store the result of every scan as an immutable snapshot keyed by repository, branch, commit SHA and "
      "timestamp. A snapshot stores **facts** about the code at that commit - findings with their evidence, the "
      "per-file risk score and churn factor, the tree, the commit SHA, the scan time, the finding count and the "
      "model version. Every **score** - priority, file debt, health score, grade, delta, category breakdown - is a "
      "function of the active profile and is therefore **derived on read**, never stored as truth, because an "
      "editable profile (FR-20) would leave a stored score stale or force an update that breaks append-only "
      "immutability. Denormalised score columns are permitted only as a cache stamped with the profile that "
      "produced them. Services remain stateless."),
     ("FR-22", "v1.0", "Account menu",
      "Account control at the foot of the left rail exposing sign out, with settings and billing stubbed."),
     ("FR-23", "v2", "Team and RBAC",
      "Multi-user workspaces with org-admin / manager / developer / viewer roles, invitations and auto-shared "
      "projects; repository access uses the workspace's GitHub App installation."),
     ("FR-24", "v1.0", "Critical-security visibility floor",
      "Critical security findings are never suppressed or down-weighted below visibility, regardless of the "
      "active profile or accepted-debt suppressions. Three independent mechanisms enforce this: (1) severity is "
      "not user-settable, so hardcoded-secret = critical is fixed in the rule register; (2) the security category "
      "is excluded from the trust slider, so source_trust is always 1.0 for security findings and no position of s "
      "can de-weight them; and (3) critical security findings are pinned into the visible list regardless of "
      "computed priority, even at the minimum permitted security weight of 0.1. Mechanism 3 shall be implemented "
      "as code now that FR-20 exposes weight sliders."),
     ("FR-25", "supporting", "ML evaluation versus rule baseline",
      "The SATD classifier and risk model are evaluated and documented with precision, recall and F1 (and AUC "
      "for risk) against the deterministic rule baseline, on held-out data and on real repositories. Per-class "
      "reporting with support counts is mandatory for the SATD classifier: documentation (54 instances) and test "
      "(85) are two orders of magnitude rarer than code-design (2,703), so an averaged figure would conceal "
      "near-total failure on the smallest classes. Accuracy is never quoted. The source paper's F1 of 0.611 "
      "covers a four-type task over four sources and is context, not a baseline; the baseline is the rule "
      "engine.")],
    [700, 1150, 2300, 5210], "Table 3-1. Functional requirement register (v1.0 and later releases).")

h3("FR-1 [v1.0] Authentication and session")
body_text("Trigger: an unauthenticated visitor opens the application and is presented with the Login page. "
          "Behaviour: the system offers a single \"Sign in with GitHub\" action; on success it establishes an "
          "authenticated session and lands the user on the Projects destination. Input: the GitHub OAuth "
          "authorization code. Output: an authenticated session bound to a workspace, and the user's project "
          "list. A signed-in user can sign out from the Account menu at the foot of the left rail. Session "
          "state is carried per request as a token so that application services remain stateless.")

h3("FR-6 [v1.0] Repository scan (manual, asynchronous, cancellable)")
body_text("Trigger: the user activates the Scan control in the dashboard top navigation for the active project "
          "and selected branch. Behaviour: the backend records a scan, enqueues an asynchronous job and returns "
          "a scan identifier immediately, so the user interface never blocks. The control behaves as a state "
          "machine - idle shows \"Scan\"; while running it shows \"Scanning NN%\" with a Stop control; on "
          "completion it returns to idle and the dashboard reflects the new snapshot; on cancel the worker stops "
          "and the previous snapshot is left intact; on error it returns to idle with an error indication. If the "
          "branch head commit SHA equals the last scanned SHA, the system may skip re-analysis and reuse the "
          "stored snapshot. In v1.0 scans run only on user request. Inputs: project identifier, branch name. "
          "Outputs: a scan identifier, progress updates, and on success a new immutable snapshot.")

h3("FR-8.1 [v1.0] Severity and category are assigned at detection (normative)")
body_text("Trigger: a detector produces a finding. Behaviour: every rule definition carries a fixed category and "
          "a fixed severity alongside its message template. The rule knows what it detected, so it knows how bad "
          "it is; both values are written onto the finding at detection time and are never recomputed - not by "
          "scoring, not by a profile change, not by the client. The table below is the assignment for the v1.0 "
          "rule set; the complete register, including the message template for each rule, is Appendix C.")
table(["Rule id", "Trigger", "Category", "Severity"],
      [("complex-function", "CCN > 15", "code-design", "Medium"),
       ("long-method", "function > 80 NLOC", "code-design", "Medium"),
       ("deep-nesting", "nesting > 4", "code-design", "Medium"),
       ("duplicate-block", "duplicated block detected", "code-design", "Low"),
       ("large-file", "file > 800 NLOC", "code-design", "Low"),
       ("hardcoded-secret", "regex and entropy on a high-entropy value assigned to a key, token or secret name",
        "security", "Critical"),
       ("sql-concat", "SQL string concatenation", "security", "High"),
       ("dangerous-eval", "eval or exec usage", "security", "High")],
      [1900, 3660, 1900, 1900], "Table 3-1a. Severity and category assigned by each v1.0 rule.")
body_text("The rule engine and the security patterns assign severity from the table above. The SATD classifier "
          "(FR-9) predicts category only and has no severity output, so every SATD finding is assigned a severity "
          "of medium in v1.0. The risk model (FR-10) assigns neither category nor severity, because it produces "
          "no list row at all. Consequently no machine-learning model assigns a severity in v1.0: severity is "
          "fully deterministic, which is what makes the critical-security visibility floor (FR-24) and "
          "configurable prioritization (FR-20) safe to defend.",
          bold_lead="Assignment by source. ")
body_text("A rule emits the same severity regardless of how far the measured value exceeds its threshold - "
          "complex-function is Medium whether the cyclomatic complexity is 16 or 45 - and a worse file simply "
          "accumulates more findings. Graduating severity by the size of the exceedance is a [v1.1] refinement "
          "that changes no architecture: the rule still decides, still at detection time.",
          bold_lead="Severity is flat in v1.0. ")
body_text("The stored severity is read by exactly two consumers - scoring, which maps it to base points (FR-11), "
          "and the Refactor-First badge (FR-15) - so the ranking a user sees and the badge they see can never "
          "disagree. Neither consumer writes the value.", bold_lead="One value, read twice. ")

h3("FR-11 [v1.0] Scoring and prioritization")
body_text("Each factor answers exactly one question and has exactly one owner, and nothing is counted twice:")
table(["Factor", "Question it answers", "Owner", "Range"],
      [("base_points(severity)", "How bad is this kind of problem?", "System - the rule register (FR-8.1) or the "
                                                                     "marker table (FR-9.2)", "1 / 3 / 5 / 8"),
       ("category_weight[category]", "How much does this team care about this type of debt?", "User - 5 sliders (FR-20)",
        "0.1 - 3.0"),
       ("source_trust(finding)", "How much does this team trust the rules versus the model?", "User - trust slider s "
                                                                                             "(FR-20)", "0.5 - 1.5; "
                                                                                                        "always 1.0 for "
                                                                                                        "security"),
       ("churn_factor(file)", "How actively is this file being changed?", "Evidence - PyDriller", "1.0 - 2.0"),
       ("risk_factor(file)", "How bug-prone is this file?", "Model - ML-2 (FR-10)", "1.0 - 2.5")],
      [2200, 3400, 2400, 1360], "Table 3-1a. The five factors of finding_priority.")
body_text("Trigger: the completion of a scan, or a change of the active scoring profile. Behaviour: the system "
          "fuses findings into scores with a pure function that reads the active profile (six per-category "
          "weights and the trust slider s) and any accepted-debt suppressions. Because suppressions and weights "
          "are applied only at scoring time, changing a profile or accepting a finding never requires a re-scan. "
          "The system computes finding priority, per-file debt, the repository health score and its grade as "
          "defined in Table 3-1, with severity base points Critical 8, High 5, Medium 3 and Low 1. The 90-day "
          "churn window is measured backwards from the committer date of the scanned commit, so re-scanning the "
          "same SHA always reproduces the same score. Outputs: per-finding priority, per-file debt scores, the "
          "repository health score, grade, delta and category breakdown, all persisted in the snapshot.")
body_text("The risk score enters scoring only through risk_factor, in the same shape as churn_factor - a bounded "
          "per-file multiplier applied to every finding in that file. It contributes no additive term to file_debt, "
          "so it is never counted twice, and a file with no findings accrues no debt from risk alone. A multiplier "
          "is also the correct shape: it scales proportionally, so a Critical finding in a fragile file gains more "
          "than a Low one, whereas an additive term would shift both equally.",
          bold_lead="Risk multiplies; it does not add. ")
info("[TO FILL - k shall be recalibrated on the golden repositories. file_debt changed scale under CR-001 (an "
     "additive term was removed and a multiplier of up to 2.5x introduced), so any previously chosen value is "
     "invalid. Do not quote a health score until this is done.]")

info("[TO FILL - expand FR-2 to FR-5 and FR-7 to FR-25 into their own sub-sections using the pattern above "
     "(trigger, behaviour, inputs, outputs). Add any v1.0 requirement the team wants that is not yet captured, "
     "such as per-screen error and empty-state behaviour, and give it a stable FR identifier.]")

# ---------- 3.2 Usability ----------
h2("Usability")
info("[Guidance: this section includes all requirements that affect usability - required training time for a "
     "normal and a power user, measurable task times for typical tasks, and conformance to recognised usability "
     "standards.]")
body_text("Usability is this product's core differentiator: competing tools are technically capable but "
          "overwhelm small teams with unranked findings and heavy configuration. The requirements below make "
          "\"low noise\" testable rather than aspirational. Each requirement states a measurable target and the "
          "method by which it will be verified.")

h3("Learnability and training time")
table(["ID", "Requirement", "Measurable target", "Verification"],
      [("U-1", "[v1.0] A first-time developer shall complete the primary journey - sign in, connect a public "
               "repository, run a scan and open the top-priority finding - unaided, with no training and no "
               "documentation.",
        "At least 4 of 5 test participants succeed unaided; median completion time <= 5 minutes.",
        "Moderated usability test with 5 participants."),
       ("U-2", "[v1.0] No formal training shall be required for any user class.",
        "0 hours of formal training; self-guided orientation for a developer <= 10 minutes; a "
        "viewer/stakeholder requires none.",
        "Same usability-test session.")],
      W4, "Table 3-2. Learnability and training-time requirements.")
info("[TO FILL - confirm the participant count and the time targets with the team, and record whether the "
     "usability test runs before submission or during the September testing phase. Do not delete the numbers: "
     "a target that has not been measured yet is still a requirement.]")

h3("Task efficiency and low-noise presentation")
table(["ID", "Requirement", "Measurable target", "Verification"],
      [("U-3", "[v1.0] On an already-scanned project, each core task shall complete within a stated interaction "
               "budget measured from dashboard load.",
        "Identify the highest-priority issue <= 30 s and <= 2 clicks; open finding detail 1 click; filter by "
        "debt type <= 2 interactions; switch branch <= 2; load a past snapshot <= 2.",
        "Automated end-to-end tests assert the click paths."),
       ("U-4", "[v1.0] The dashboard shall present findings ranked by priority and shall never present an "
               "unranked dump of raw results.",
        "The initial Refactor-First view shows at most 20 rows, sorted by priority descending; everything else "
        "is reached by explicit filtering or drill-in.",
        "Component test on the sort and slice."),
       ("U-5", "[v1.0] A non-technical stakeholder shall correctly interpret the health score, grade, delta "
               "direction and trend direction without assistance.",
        "At least 4 of 5 non-technical participants interpret all four correctly.",
        "Comprehension questions in the usability test."),
       ("U-6", "[v1.0] Every finding shall be explained in exactly one plain-English sentence naming the symbol, "
               "the measured value, the threshold crossed and the corrective action, with no unexplained jargon.",
        "100% of findings; <= 140 characters per reason.",
        "Review of the reason-template table (Appendix C).")],
      W4, "Table 3-3. Task-efficiency and low-noise requirements.")

h3("Accessibility and standards conformance")
table(["ID", "Requirement", "Measurable target", "Verification"],
      [("U-7", "[v1.0] The user interface shall conform to WCAG 2.1 Level AA for contrast, including every "
               "severity badge and every heat-map tint.",
        "Contrast ratio >= 4.5:1 for body text and >= 3:1 for user-interface components and graphical objects; "
        "0 violations.",
        "Automated accessibility audit (axe-core / Lighthouse)."),
       ("U-8", "[v1.0] Colour shall never be the sole carrier of meaning.",
        "Heat-map tint is always accompanied by a numeric health score and letter grade; severity is always a "
        "text-labelled badge, never a coloured dot alone; 0 colour-only signals.",
        "Manual audit plus greyscale review."),
       ("U-9", "[v1.0] Every interactive control shall be reachable and operable by keyboard, in a logical tab "
               "order, with a persistently visible focus indicator.",
        "100% of controls: rail links, branch selector, Scan and Stop, debt-type filter, finding rows, "
        "file-tree nodes, detail-panel close.",
        "Keyboard-only walkthrough plus automated audit."),
       ("U-15", "[v1.1] Accessibility hardening: the file tree adopts the ARIA tree pattern with arrow-key "
                "navigation, and scan progress is announced through a live region.",
        "Deferred - not required for v1.0 acceptance.",
        "Deferred to the v1.1 accessibility pass.")],
      W4, "Table 3-4. Accessibility and standards-conformance requirements.")

h3("Feedback, error recovery and consistency")
table(["ID", "Requirement", "Measurable target", "Verification"],
      [("U-10", "[v1.0] Any action that takes longer than one second shall show progress within that second.",
        "Scans surface a live state machine (idle, queued, running NN%, done or error) with a Stop control and a "
        "terminal notification; data loads show skeleton placeholders rather than blank regions; no "
        "unacknowledged action exceeds 1 s.",
        "End-to-end test of the scan state machine."),
       ("U-11", "[v1.0] Error recovery shall be non-destructive: a failed or cancelled scan leaves the previous "
                "snapshot displayed and usable, and every error message states what failed and the next action.",
        "0 dead-end error states.",
        "End-to-end cancel and failure paths (see R-1)."),
       ("U-12", "[v1.0] Presentation shall be consistent: one design system, one rendering per severity and "
                "grade on every screen, and user-interface copy that uses the exact terms defined in Section 1.3.",
        "0 divergent design tokens; all colours resolved from CSS variables; no synonyms for \"finding\", "
        "\"snapshot\" or \"debt category\".",
        "Code-review gate (see DC-2)."),
       ("U-13", "[v1.0] The application shall be fully usable from a laptop viewport upwards, with no horizontal "
                "scrolling of the dashboard.",
        "Minimum supported viewport 1280 x 720; the two-column dashboard collapses to a single column below the "
        "large breakpoint.",
        "Automated viewport test."),
       ("U-14", "[v1.0] Every list shall have an empty state naming the action that fills it.",
        "100% of list views (for example \"No repositories yet - connect one to see its health.\").",
        "Screen-by-screen review.")],
      W4, "Table 3-5. Feedback, error-recovery and consistency requirements.")

# ---------- 3.3 Reliability ----------
h2("Reliability")
info("[Guidance: specify availability (percentage of time available, hours of use, degraded modes), Mean Time "
     "Between Failures, Mean Time To Repair, accuracy, and the maximum acceptable defect rate by severity.]")
table(["ID", "Requirement", "Target"],
      [("R-1", "[v1.0] A failed or cancelled scan shall not corrupt the previous snapshot; the dashboard "
               "continues to show the last good snapshot.", "100% of failure and cancel paths."),
       ("R-2", "[v1.0] Detection shall be deterministic and reproducible: the rule engine yields identical "
               "findings for identical inputs, and re-scanning the same commit SHA reproduces the same score.",
        "100% reproducible."),
       ("R-3", "[v1.0] Machine-learning outputs shall be presented as risk and health indicators with documented "
               "error rates, never as guarantees.", "Metrics published with the release (see FR-25)."),
       ("R-4", "[v1.0] Degraded mode: if the machine-learning service is unavailable, the scan shall complete "
               "using the rule engine alone and the result shall be marked as degraded.",
        "Scan still completes; no data loss."),
       ("R-5", "[v1.0] Availability of the hosted service, measured monthly and excluding announced maintenance "
               "windows.", "99.0% (proposed)"),
       ("R-6", "[v1.0] Mean Time Between Failures and Mean Time To Repair for the hosted service.",
        "MTBF >= 200 hours; MTTR <= 4 hours (proposed)"),
       ("R-7", "[v1.0] Accuracy of the machine-learning components, reported per FR-25 and never presented as a "
               "guarantee.",
        "SATD classifier F1 >= 0.75 on held-out comments; risk model AUC >= 0.70 and F1 reported alongside "
        "precision and recall (proposed - defect prediction is a hard problem and these figures are consistent "
        "with published baselines; confirm after the first evaluation run)"),
       ("R-8", "[v1.0] Maximum defect rate at submission.",
        "0 open critical defects; 0 open high defects; at most 5 open medium defects (proposed)")],
      [700, 5000, 3660], "Table 3-6. Reliability requirements.")
info("[Confirm the four figures marked (proposed) with the team and delete the marker. Reliability is assessed on "
     "whether targets are stated and justified, not on how ambitious they are - so keep the numbers even if the "
     "team lowers them.]")

# ---------- 3.4 Performance and Security ----------
h2("Performance and Security")
info("[Guidance: include specific response times (average and maximum), throughput, capacity, degradation modes "
     "and resource utilisation. Security controls follow in 3.4.2.]")

h3("Performance")
table(["ID", "Requirement", "Target"],
      [("P-1", "[v1.0] The dashboard shall remain responsive during analysis: scans run on asynchronous workers "
               "off the request path and the user interface polls for progress.",
        "No blocking request longer than 1 s during a scan."),
       ("P-2", "[v1.0] Reads of a stored snapshot (health, trend, list, tree) are pure database reads with no "
               "computation on read.", "95th-percentile server response < 500 ms (proposed)"),
       ("P-3", "[v1.0] Profile switches (including dragging a weight slider) and folder drill-ins recompute from "
               "stored findings without re-scanning.", "< 200 ms perceived (proposed)"),
       ("P-4", "[v1.0] Dashboard first meaningful paint on a broadband connection.", "< 2 s (proposed)"),
       ("P-5", "[v1.0] Scan throughput.",
        "A 50-KLOC repository completes within 5 minutes; each worker sustains 4 concurrent scans (proposed - "
        "confirm against the first end-to-end pipeline run)"),
       ("P-6", "[v1.0] Capacity.",
        "Scaled for 50 teams of 5 developers per the Feasibility Report; assume 1 repository per team, a scan per "
        "working day, and a stored snapshot of the order of 1 MB, giving roughly 1 GB of snapshot storage per year "
        "(proposed)"),
       ("P-7", "[v1.0] Degradation mode: under worker saturation, scans queue rather than fail, and the user "
               "interface reports the queued state.", "Queue depth surfaced; no dropped scans.")],
      [700, 5000, 3660], "Table 3-7. Performance requirements.")

h3("Security")
table(["ID", "Requirement"],
      [("S-1", "[v1.0] Repository access shall be least-privilege and read-only; the user explicitly selects "
               "which repositories the system may access, and the system cannot read any other repository."),
       ("S-2", "[v1.0] All client-server communication shall use HTTPS; tokens and secrets are stored "
               "server-side and are never exposed to the client."),
       ("S-3", "[v1.0 seam] Tenants shall be logically isolated using PostgreSQL Row-Level Security keyed on "
               "workspace_id; the tenant is always derived from the session server-side and is never accepted "
               "from the client."),
       ("S-4", "[v1.1] Private-repository access shall use GitHub App installation tokens, which are "
               "fine-grained and revocable."),
       ("S-5", "[v1.0] The system shall not retain repository source code beyond what the analysis requires; "
               "clones are working copies and are removed after the scan."),
       ("S-6", "[v1.0] Authentication is delegated to GitHub OAuth; the system stores no user passwords.")],
      [700, 8660], "Table 3-8. Security requirements.")

# ---------- 3.5 Supportability ----------
h2("Supportability")
info("[Guidance: this section indicates any requirements that enhance the supportability or maintainability of "
     "the system - coding standards, naming conventions, maintenance utilities, configurability, testability "
     "and diagnostics.]")
body_text("These requirements answer a single question: when a new maintainer joins this project in six months, "
          "what has been built in so that they can change the system safely?")

h3("Coding standards and conventions")
table(["ID", "Requirement"],
      [("SP-1", "[v1.0] The frontend shall be written in TypeScript with strict mode enabled and shall pass "
                "ESLint and Prettier; the backend and machine-learning code shall follow PEP 8, enforced by an "
                "automated formatter and linter, with type hints on all public functions. TO FILL - confirm the "
                "Python toolchain (for example ruff and black)."),
       ("SP-2", "[v1.0] Commit messages shall follow the Conventional Commits convention and branches shall be "
                "named by type (feat/, fix/, docs/, test/), so that history is machine-readable and release "
                "notes are derivable."),
       ("SP-3", "[v1.0] Every change shall reach the main branch through a reviewed pull request; linting, type "
                "checking, unit tests and end-to-end tests are the merge gate.")],
      [700, 8660], "Table 3-9. Coding standards and conventions.")

h3("Modularity, configurability and extensibility")
table(["ID", "Requirement"],
      [("SP-4", "[v1.0] A single data contract shall be the source of truth for every shape crossing the "
                "frontend, backend and database boundary. It changes only through a pull request reviewed by "
                "both sides; a field that is not in the contract does not exist."),
       ("SP-5", "[v1.0] The codebase shall be layered (presentation, application/domain, data) within a "
                "monorepo. The presentation layer never accesses the database directly; layer-skipping is a "
                "review rejection."),
       ("SP-6", "[v1.0] Volatile third-party libraries shall sit behind exactly one boundary so that replacing "
                "one is a single-file change: all charts through one chart wrapper, the file-tree implementation "
                "behind one component, all HTTP through one API client module."),
       ("SP-7", "[v1.0] Scoring shall remain a pure function over stored findings, so that changing a weight "
                "profile or a suppression re-scores instantly and never requires a re-scan or a data migration."),
       ("SP-8", "[v1.0] Rule thresholds, severity base points and profile weights shall be configuration or "
                "data, never literals in code, so that calibration is a configuration change rather than a "
                "release."),
       ("SP-9", "[v1.0] All environment-specific values shall be supplied as environment variables (API base "
                "URL, mocking switch, database and broker URLs, GitHub credentials). No secret is committed to "
                "the repository and no environment difference requires a code edit."),
       ("SP-17", "[v1.0] The system shall be extensible without engine changes: adding a rule is one threshold "
                 "entry plus one reason template; adding a detector is a new source value; adding a language is "
                 "a per-language rule pack plus recalibration.")],
      [700, 8660], "Table 3-10. Modularity, configurability and extensibility.")

h3("Testability and diagnosability")
table(["ID", "Requirement"],
      [("SP-10", "[v1.0] The frontend shall be fully testable without a backend: the same mock handlers serve "
                 "development, unit tests and end-to-end tests. Every data hook and every dashboard component "
                 "has a unit test, and the core journeys (dashboard render, scan lifecycle, branch switch) have "
                 "end-to-end tests."),
       ("SP-11", "[v1.0] The rule engine shall be deterministic, so that regression tests are exact rather than "
                 "statistical."),
       ("SP-12", "[v1.0] Logs shall be structured and correlated by scan identifier across the API, the broker, "
                 "the worker and the machine-learning service, so that one scan is traceable end to end."),
       ("SP-13", "[v1.0] Every scan shall persist its terminal state (phase and error) so that a user-reported "
                 "failure is diagnosable from the database without inspecting logs.")],
      [700, 8660], "Table 3-11. Testability and diagnosability.")

h3("Machine-learning, documentation and release maintainability")
table(["ID", "Requirement"],
      [("SP-14", "[v1.0] Models shall be trained offline and loaded at runtime as versioned artifacts, so that "
                 "replacing a model requires no application change."),
       ("SP-15", "[v1.0] Each stored snapshot shall record the model version and the scoring profile that "
                 "produced it. Without this, re-training silently makes historical scores incomparable and the "
                 "trend chart loses meaning. TO FILL - confirm with the backend that the scan record carries a "
                 "model_version column."),
       ("SP-16", "[v1.0] The risk model shall support periodic re-training on recent data with no architectural "
                 "change."),
       ("SP-18", "[v1.0] The SRS, the Software Architecture Document, the backend analysis-engine document and "
                 "per-application README files shall be maintained alongside the code, together with decision "
                 "records for choices that are expensive to reverse, each stating the decision, its date and "
                 "its rationale."),
       ("SP-19", "[v1.0] The database schema shall evolve only through versioned, additive migrations, and enum "
                 "string values shall be frozen before the first migration because they become both database "
                 "constraints and user-interface values."),
       ("SP-20", "[v1.0] Every component shall ship as an independently buildable and independently scalable "
                 "container, and builds shall be reproducible from committed lock files.")],
      [700, 8660], "Table 3-12. Machine-learning, documentation and release maintainability.")

# ---------- 3.6 Design Constraints ----------
h2("Design Constraints")
info("[Guidance: design constraints represent design decisions that have been mandated and must be adhered to - "
     "languages, tools, architecture, standards compliance and hardware limitations.]")
table(["ID", "Constraint"],
      [("DC-1", "The technology stack is fixed: Next.js (App Router) with TypeScript and Tailwind/shadcn on the "
                "frontend; FastAPI with Celery and Redis on the backend; Python with scikit-learn, Lizard and "
                "PyDriller for analysis and machine learning; PostgreSQL for data; containerized with Docker."),
       ("DC-2", "Charts are rendered through a single chart wrapper; colours come from CSS variables and are "
                "never hardcoded; all displayed data comes from the data contract and is never hardcoded."),
       ("DC-3", "Machine learning uses supervised models trained once and then used for inference. No "
                "reinforcement learning. Per-repository calibration and feedback learning are deferred."),
       ("DC-4", "Language support is agnostic by architecture but validated on Python and JavaScript/TypeScript "
                "only; new languages require per-language security rules and recalibration."),
       ("DC-5", "Scope is locked to the Project Proposal; new ideas are recorded on the future-improvements "
                "list rather than added to v1.0."),
       ("DC-6", "Only free and open-source tools, datasets and hosting may be used; the project budget is zero.")],
      [700, 8660], "Table 3-13. Design constraints.")

# ---------- 3.7 Documentation ----------
h2("On-line User Documentation and Help System Requirements")
info("[Guidance: describe the requirements for on-line user documentation, help systems and help-about notices.]")
table(["ID", "Requirement"],
      [("DOC-1", "[v1.0] The application shall include in-context help: tooltips and short explanations for the "
                 "health score, the grade bands, severity levels and the meaning of each debt category."),
       ("DOC-2", "[v1.0] Every empty state shall explain what the screen will show and name the action that "
                 "fills it (see U-14)."),
       ("DOC-3", "[v1.0] The repository shall carry a README describing setup, environment variables and the "
                 "test commands, sufficient for a new developer to run the system locally."),
       ("DOC-4", "[v1.1] A dedicated help page or first-run onboarding tour. TO FILL - confirm whether the team "
                 "wants this in scope, and for which release.")],
      [900, 8460], "Table 3-14. Documentation and help requirements.")

# ---------- 3.8 Purchased components ----------
h2("Purchased Components")
info("[Guidance: describe any purchased components, their licensing and any usage restrictions.]")
body_text("No purchased or paid components are used. All tools, libraries and datasets are free or open-source, "
          "consistent with the zero-cost position of the Feasibility Report. Third-party payment processing "
          "applies only to the [v2] billing feature and is out of scope for v1.0.")
table(["Component", "Role", "Licence / cost"],
      [("Next.js, React, Tailwind CSS, shadcn/ui", "Frontend framework and user-interface components",
        "Open-source (MIT) - no cost"),
       ("FastAPI, Celery, Redis", "Backend API, task queue and message broker", "Open-source (MIT / BSD) - no cost"),
       ("Python, scikit-learn, Lizard, PyDriller", "Analysis and machine learning", "Open-source (BSD / MIT) - no cost"),
       ("PostgreSQL", "Relational database", "Open-source (PostgreSQL licence) - no cost"),
       ("Docker", "Containerization", "Free tier - no cost"),
       ("SATD dataset (Li et al.)", "Training data for the SATD classifier", "MIT (c) 2022 Yikun Li - permissive, commercial use permitted, attribution required"),
       ("GHPR defect dataset", "Training data for the risk model", "TO FILL - confirm the research licence"),
       ("Stripe [v2]", "Billing for the Team tier", "Commercial, transaction fee - not used in v1.0")],
      [3000, 3400, 2960], "Table 3-15. Third-party components and their licences.")
body_text("The SATD dataset used by v1.0 is MIT-licensed (c) 2022 Yikun Li - permissive, commercial use "
          "permitted, attribution required - so it places no restriction on this project. The separate Technical "
          "Debt Dataset is CC BY-NC-SA 4.0 and therefore research and non-commercial only; v1.0 does not use it.",
          bold_lead="Dataset licences. ")
info("[TO FILL - confirm the GHPR defect-dataset licence before ML-2 training begins.]")

# ---------- 3.9 Interfaces ----------
h2("Interfaces")
info("[Guidance: describe the interfaces the application must support, without screenshots. For user "
     "interfaces, list the functionality and the required menu items, panels, text boxes, buttons and drop-down "
     "lists; for software interfaces, the interfaces to servers, web services and libraries; for hardware "
     "interfaces, the client-side prerequisites; for communication interfaces, the protocols used. Include a "
     "draft block diagram of the main interfaces.]")

h3("User Interfaces")
body_text("The graphical user interface is a single-page-style web application: a persistent left navigation "
          "rail with a content area that swaps by destination. Layout is described structurally; no screenshots "
          "are included at this stage.")
table(["Element", "Required contents"],
      [("Left navigation rail", "Product name; navigation items Projects, Dashboard, Scan History, Profiles and "
                                "Team (badged as v2); collapsible to icons; the active item is highlighted; an "
                                "Account control pinned at the foot."),
       ("Account menu", "Sign out; Settings and Billing present but stubbed for v1.0."),
       ("Notifications", "Transient toast messages for action outcomes: scan complete, scan failed, scan cancelled.")],
      W2, "Table 3-16. Global user-interface elements present on every authenticated screen.")

table(["Screen", "Required controls and fields", "States"],
      [("Login", "A single prominent \"Sign in with GitHub\" button, with the product name and tagline. No local "
                 "user name or password field exists.", "Idle; authenticating; authentication error."),
       ("Projects", "Connect panel: a labelled repository-URL text box and a Connect button with inline "
                    "validation. Project list (vertical): per row the repository name, owner, visibility badge, a "
                    "health hint (score, grade, delta) where a snapshot exists, and a Select action.",
        "Loading; empty (\"No repositories yet - connect one to see its health.\"); populated; invalid-URL error."),
       ("Dashboard - dashboard mode", "Top navigation: active repository name, branch drop-down list, Scan control "
                                      "(Scan -> Scanning NN% with Stop -> idle), last-analysed timestamp and short "
                                      "commit SHA. Left column: Overall Health card (score 0-100, grade A-E, delta, "
                                      "critical/high count, category pie), health-trend chart, and the Refactor-First "
                                      "list with a debt-type filter drop-down whose rows show a category chip, a "
                                      "severity chip, file:line, the one-line reason and - for SATD rows only - a "
                                      "SATD source chip. Right column: the hotspot file-tree heat map with per-file "
                                      "risk badges, expand, collapse, hover and select.",
        "Loading (skeleton placeholders); loaded; load error; scanning; scan error."),
       ("Dashboard - detail mode", "Selecting a finding replaces the health card and trend chart in place with the "
                                   "finding detail (evidence, one-line reason, file:line:symbol, and the region "
                                   "reserved for the [v1.1] snippet); the file tree auto-expands and highlights that "
                                   "finding's file; the Refactor-First list condenses to a strip so the user can move "
                                   "between findings. Closing restores dashboard mode. No overlay and no blurred "
                                   "background. The selected finding appears in the URL.",
        "Entering; loaded; finding not found; leaving."),
       ("Scan History", "A table of past snapshots: date, commit SHA, health score, grade, delta and finding "
                        "count; selecting a row loads that snapshot into the dashboard read-only.",
        "Loading; empty (\"No scans yet\"); populated."),
       ("Profiles", "Three preset buttons (Balanced - the default - Security-first and Delivery-speed) that seed the "
                    "controls; six category weight sliders (security, code-design, defect, requirement, documentation, test; "
                    "clamped 0.1-3.0); one trust slider labelled at its ends \"trust the rules\" and \"trust the "
                    "model\" (0-1, default 0.5); and a Reset to preset action. Every change re-scores instantly with "
                    "no re-scan.", "Loaded; applying; modified-from-preset."),
       ("Team [v2]", "Placeholder only in v1.0; badged v2 in the navigation rail.", "Not applicable.")],
      [1300, 5400, 2660], "Table 3-17. Screen inventory: required panels, fields and states.")

body_text("Severity is always shown as a text-labelled badge, and grade and heat-map tint are always accompanied "
          "by the numeric score, so that colour is never the only carrier of meaning (U-8). All colours resolve "
          "from design tokens rather than hardcoded values (DC-2). All displayed values come from the data "
          "contract: the dashboard performs no calculation of its own. Every list defines a loading, an empty "
          "and an error state (U-14).", bold_lead="User-interface conventions (normative). ")

info("[TO FILL - insert Figure 3-1, a block diagram of the main user interfaces, drawn in draw.io or "
     "Lucidchart per the figure guidelines (black text of at least 12 pt on a white fill, numbered caption, and "
     "at least one sentence of description). The structure to draw is given below.]")
table(["Diagram level", "Nodes"],
      [("Actor", "User"),
       ("Navigation", "Left rail (persistent)"),
       ("Destinations", "Projects (connect, list, select) | Dashboard | Scan History | Profiles | Team [v2] | "
                        "Account menu (sign out)"),
       ("Dashboard panels", "Top navigation (repository, branch, Scan/Stop, last commit SHA) | Overall Health "
                            "card with category pie | Health-trend chart | Refactor-First list with debt-type "
                            "filter | Hotspot file-tree heat map"),
       ("Shared panel", "Finding-detail slide-over, opened from either the Refactor-First list or the file tree")],
      [2000, 7360], "Table 3-18. Structure to be drawn as the Figure 3-1 interface block diagram.")

h3("Hardware Interfaces")
body_text("Code Sage AI is a browser-delivered web application and interfaces with no specialised hardware: no "
          "peripherals, sensors, printers or device drivers are used. The requirements below are therefore "
          "client-side prerequisites and hosting expectations rather than hardware protocols.")
table(["Item", "Requirement"],
      [("Client device", "Any desktop or laptop computer capable of running a current web browser. No "
                         "installation and no local storage are required."),
       ("Client display", "Minimum supported viewport 1280 x 720 (see U-13)."),
       ("Client processor and memory", "A dual-core processor and 4 GB RAM (proposed). The client renders only - "
                                       "all analysis runs server-side - so requirements are modest and are "
                                       "effectively those of the browser itself."),
       ("Client storage", "None. No installation and no local data storage are required."),
       ("Client network", "A broadband internet connection; the application is unusable offline."),
       ("Server / worker (hosting)", "Commodity cloud compute running containers. Each worker requires local disk "
                                     "for a clone of the repository under analysis - provisionally at least 2 GB per "
                                     "concurrent scan, released on completion (proposed) - and outbound HTTPS access "
                                     "to the Git host.")],
      [2400, 6960], "Table 3-19. Hardware interfaces and client-side prerequisites.")

h3("Software Interfaces")
table(["Interface", "Purpose", "Direction and protocol"],
      [("GitHub OAuth", "User authentication (FR-1).", "Outbound; HTTPS, OAuth 2.0 authorization-code flow."),
       ("GitHub REST API", "Repository metadata: name, owner, visibility, default branch, branch list and head "
                           "commit SHAs.",
        "Outbound, read-only; HTTPS with JSON, using ETag-based conditional requests to mitigate rate limits."),
       ("Git (clone)", "Fetch the working tree and the commit history for analysis, consuming no API quota.",
        "Outbound, read-only; Git over HTTPS."),
       ("GitHub App [v1.1]", "Private-repository access with user-selected repositories (least privilege).",
        "Outbound; HTTPS with installation tokens."),
       ("GitLab API [v2]", "Support for a second Git host.", "Outbound; HTTPS with JSON."),
       ("Payment provider [v2]", "Team-tier billing.", "Outbound; HTTPS with JSON."),
       ("Web browser runtime", "Rendering host for the frontend. A Service Worker is registered only by the "
                               "development and test mock layer, never in production.",
        "Inbound; standard HTML, CSS and JavaScript.")],
      [1900, 4300, 3160], "Table 3-20. External software interfaces.")

table(["Interface", "Between", "Protocol and format"],
      [("Application REST API", "Frontend and backend",
        "HTTPS with JSON; payload shapes governed by the data contract; specified as an OpenAPI document in the "
        "Software Architecture / Design Document."),
       ("Task queue", "Backend API and asynchronous workers",
        "Redis protocol on a private network, carrying job enqueue and progress state."),
       ("Machine-learning inference", "Worker and machine-learning service",
        "Batched comments in, SATD label and category out; per-file numeric feature vector in, risk score 0-1 out."),
       ("Persistence", "API and workers, and the database", "SQL over TCP with TLS, using pooled connections."),
       ("Mock layer (development and test only)", "Frontend and the mock service worker",
        "Intercepts requests at the network boundary and honours the identical contract, so that moving to the "
        "real backend is a base-URL change rather than a rewrite.")],
      [2400, 2800, 4160], "Table 3-21. Internal software interfaces.")

table(["Method", "Path", "Purpose"],
      [("POST", "/api/auth/github", "Sign in; returns the user and the accessible repositories (FR-1)."),
       ("GET", "/api/projects", "List the connected projects (FR-4)."),
       ("POST", "/api/projects", "Connect a repository by URL (FR-3). TO FILL - specified here but not yet "
                                 "implemented in the mock layer; raise with the backend."),
       ("GET", "/api/repos/{repoId}/branches", "Branch list with head commit SHA and default flag (FR-5)."),
       ("GET", "/api/repos/{repoId}/health?branch=", "The full dashboard payload for one branch snapshot "
                                                     "(FR-12 to FR-18)."),
       ("GET", "/api/repos/{repoId}/scans", "Scan history: the stored snapshot summaries (FR-19)."),
       ("POST", "/api/repos/{repoId}/scan", "Start a scan for a branch; returns a scan identifier and phase (FR-6)."),
       ("GET", "/api/repos/{repoId}/scan/{scanId}", "Poll the phase and progress of a running scan (FR-6)."),
       ("POST", "/api/repos/{repoId}/scan/{scanId}/stop", "Cancel a running scan (FR-6)."),
       ("GET", "/api/profiles", "List the available scoring profiles (FR-20).")],
      [900, 3400, 5060], "Table 3-22. Application REST endpoints required for v1.0.")

body_text("The SATD dataset and the defect dataset are offline, download-once training artifacts. No dataset and "
          "no external API is contacted during a scan. Although the defect dataset's name refers to pull "
          "requests, it is training data only and does not imply that the system reads pull requests at run time "
          "(see FR-7.1).", bold_lead="Datasets. ")

h3("Communications Interfaces")
table(["Channel", "Protocol and format", "Notes"],
      [("Browser to frontend and API", "HTTPS (TLS 1.2 or later), JSON over REST",
        "All external traffic is encrypted; certificates are managed at the hosting edge."),
       ("Scan progress", "Asynchronous HTTP polling of the scan-status endpoint",
        "Polling interval 2 seconds in production (proposed; the prototype polls sub-second against the mock). No "
        "WebSockets or server-sent events in v1.0: polling is sufficient for a single progress value and keeps the "
        "services stateless. A push transport becomes an option in [v2] if silent checks are added."),
       ("Backend to Git host", "HTTPS: REST for metadata and Git over HTTPS for clones", "Read-only, least privilege."),
       ("Backend to workers", "Redis protocol", "Private network only; never exposed publicly."),
       ("Services to database", "PostgreSQL wire protocol over TLS", "Private network only."),
       ("Session and credentials", "Bearer token supplied per request",
        "Services remain stateless; tokens are never held in server memory and never exposed to the client (S-2)."),
       ("Electronic mail [v2]", "SMTP", "Used only for workspace invitations in the Team tier; not used in v1.0.")],
      [2200, 3000, 4160], "Table 3-23. Communications interfaces.")

# ---------- 3.10 Database ----------
h2("Database Requirements")
info("[Guidance: define the database requirements for the system. This section states what the database must "
     "guarantee; the entity-relationship diagram and the column-level schema belong to the Data View of the "
     "Software Architecture / Design Document and are referenced from here.]")
table(["ID", "Requirement"],
      [("DB-1", "[v1.0] PostgreSQL shall store workspaces, users and memberships, repository and branch "
                "metadata, immutable scan snapshots, findings, per-file risk scores, scoring profiles and "
                "accepted-debt suppressions."),
       ("DB-2", "[v1.0] Multi-tenant isolation shall be enforced by Row-Level Security keyed on workspace_id, "
                "which shall be present on every tenant-scoped table from the first migration even while the "
                "workspace holds a single member."),
       ("DB-3", "[v1.0] Scan records shall be append-only and keyed by repository, branch, commit SHA and "
                "timestamp, so that the trend chart, scan history, delta and skip-if-unchanged behaviour are all "
                "pure reads over stored rows. A stored snapshot is never an input to a later scan, and a profile "
                "change never inserts a snapshot row (FR-20)."),
       ("DB-4", "[v1.0] Stored shapes shall satisfy the data contract, and enum values (severity, source, "
                "category, grade, finding status) shall be frozen before the first migration because they become "
                "database constraints and user-interface values simultaneously."),
       ("DB-5", "[v1.0] Indexes shall support the dashboard's read patterns, in particular lookups by "
                "repository, branch and scan time."),
       ("DB-6", "[v1.0] Backup and retention: a daily full backup with point-in-time recovery over a 7-day window; "
                "scan snapshots are retained for the life of the project and are never pruned, because the trend "
                "chart and scan history are queries over that history (proposed - confirm against the hosting "
                "provider's free-tier limits)."),
       ("DB-7", "[v1.0] The scoring profile is stored per workspace as six category weights plus the trust scalar "
                "s; it shall contain no severity column, since severity is written by the detector and must not be "
                "reachable from a user setting (FR-8.1, FR-24)."),
       ("DB-8", "[v1.0] The database shall store **facts, not scores** (FR-21). Findings with their evidence, the "
                "per-file risk score and churn factor, the commit SHA, scan time, finding count and model version "
                "are persisted; priority, file debt, health score, grade, delta and the category breakdown are "
                "derived on read under the active profile. Any denormalised score column is a cache: it shall be "
                "stamped with the profile that produced it and recomputed when the active profile differs. "
                "Optional accelerator - per snapshot, two sums per (category, source) group, being the sum of "
                "base x churn and the sum of base x churn x risk, allow an exact re-score of an entire history in "
                "a few hundred operations, because the profile factors are constant within a group.")],
      [700, 8660], "Table 3-24. Database requirements.")
info("[Scope note: the schema models the v1.0 feature set plus the multi-tenancy seam only. Structures that "
     "exist solely for later releases - invitations, billing, cross-repository dependency data and webhook state "
     "- are deliberately not modelled. The suppression table and the extended finding-status values are created "
     "in v1.0 but remain unused until [v1.1]. TO FILL - reference the exact figure number of the "
     "entity-relationship diagram in the Software Architecture Document.]")

# ---------- 3.11 Legal ----------
h2("Licensing, Legal, Copyright, and Other Notices")
info("[Guidance: define any licensing enforcement or usage-restriction requirements, together with legal "
     "disclaimers, warranties, copyright and trademark notices.]")
table(["ID", "Requirement"],
      [("L-1", "All development tools and libraries shall be used under their open-source licences; no "
               "copyright is infringed."),
       ("L-2", "Repository data shall be processed under least-privilege, user-granted access and handled in "
               "line with privacy and data-protection norms; no data is shared with third parties without "
               "authorisation."),
       ("L-3", "Public research datasets shall be used only within the terms of their licences. The Li et al. "
               "SATD dataset used by v1.0 is MIT-licensed (c) 2022 Yikun Li: permissive, commercial use "
               "permitted, attribution required. A non-commercial licence - such as the CC BY-NC-SA 4.0 of the "
               "separate Technical Debt Dataset, which v1.0 does not use - would restrict the project to "
               "research and academic use."),
       ("L-4", "The user interface shall carry the applicable copyright notice and, for the hosted service, a "
               "link to the privacy policy. TO FILL - confirm the notice wording and publish a privacy policy "
               "before any public deployment.")],
      [700, 8660], "Table 3-25. Licensing and legal requirements.")

# ---------- 3.12 Standards ----------
h2("Applicable Standards")
info("[Guidance: describe by reference any applicable standard and the specific sections that apply.]")
table(["ID", "Standard and applicability"],
      [("ST-1", "IEEE Std 830-1998, Recommended Practice for Software Requirements Specifications - the "
                "structure and content of this document."),
       ("ST-2", "WCAG 2.1 Level AA - contrast, keyboard operability and colour independence of the user "
                "interface (see U-7 to U-9)."),
       ("ST-3", "OAuth 2.0 and GitHub App least-privilege guidance - external repository access and "
                "authentication (see S-1, S-4)."),
       ("ST-4", "PostgreSQL Row-Level Security practice - multi-tenant isolation (see S-3, DB-2)."),
       ("ST-5", "Team coding conventions - PEP 8 for Python and the agreed ESLint and Prettier configuration "
                "for TypeScript (see SP-1).")],
      [700, 8660], "Table 3-26. Applicable standards.")

# ═════════════════════════ 4. SUPPORTING INFORMATION ═════════════════════════
h1("Supporting Information")
info("[Guidance: supporting information makes the SRS easier to use - table of contents, index and appendices. "
     "Appendices may include user-interface prototypes; where appendices are included the SRS should state "
     "whether they are part of the requirements.]")
body_text("The table of contents appears at the front of this document. The appendices below are part of the "
          "requirements where they define data shapes or traceability, and are informative where they record "
          "supporting material.")

h2("Appendix A - Requirements traceability")
body_text("This matrix is what demonstrates that every proposal objective is realised and that every requirement "
          "is testable.")
table(["Requirement", "Proposal objective", "Release", "Verification / test case"],
      [("FR-1 Authentication and session", "Objective 3 - centralized interactive dashboard", "v1.0",
        "TC-01 sign-in end-to-end test"),
       ("FR-2 Workspace and tenant isolation", "Objective 2 - secure multi-tenant architecture", "v1.0 seam",
        "TC-02 row-level-security isolation test"),
       ("FR-3 Connect a repository", "Objective 3", "v1.0", "TC-03 connect valid and invalid URL"),
       ("FR-6 Repository scan", "Objective 3", "v1.0", "TC-06 scan lifecycle end-to-end test (start, progress, "
                                                       "cancel, error, skip-if-unchanged)"),
       ("FR-7.1 Extraction boundary", "Objective 1 - analyse the impact of technical debt", "v1.0",
        "TC-07 same-SHA re-scan reproduces an identical snapshot"),
       ("FR-8 / FR-8.1 Rule engine and severity register", "Objective 5 - evaluate ML against the rule baseline",
        "v1.0", "TC-08 deterministic rule fixtures; register review"),
       ("FR-9 / FR-9.2 SATD classifier and severity", "Objective 4 - NLP classifier for self-admitted debt", "v1.0",
        "TC-09 held-out comment evaluation; marker-table unit tests"),
       ("FR-10 Risk model", "Objective 5", "v1.0", "TC-10 precision / recall / F1 / AUC report"),
       ("FR-11 Scoring and prioritization", "Objective 3", "v1.0",
        "TC-11 worked-example fixture; bound check that no low finding outranks a critical one in the same category"),
       ("FR-12 to FR-19 Dashboard outputs", "Objective 3", "v1.0", "TC-12 dashboard render and drill-in E2E"),
       ("FR-20 Scoring profiles", "Objective 3", "v1.0",
        "TC-20 preset seeds sliders, weights clamp, re-score without a scan"),
       ("FR-21 Snapshot persistence", "Objective 1", "v1.0", "TC-21 append-only insert; trend and history reads"),
       ("FR-24 Visibility floor", "Objective 3", "v1.0",
        "TC-24 critical security finding stays visible at the minimum security weight"),
       ("FR-25 ML evaluation", "Objective 5", "supporting", "Model evaluation report, September testing phase")],
      [2600, 3000, 1200, 2560], "Table 4-1. Requirements traceability matrix.")
info("[TO FILL - confirm the test-case identifiers against the Test Plan once it is written, and add rows for any "
     "requirement the team adds to Section 3.1.]")

h2("Appendix B - Dashboard output definitions")
body_text("The six dashboard outputs - the overall health card, the category breakdown, the health trend, the "
          "Refactor-First list, the finding-detail panel and the hotspot file tree - are defined by the shared "
          "data contract, which is the single normative source for their field names, types and permitted "
          "values. Any change to those shapes is a change to this appendix.")

h2("Appendix C - Rule register (rule id, severity, category, message template)")
body_text("This appendix is the single normative source for two requirements at once: FR-8.1, which states that "
          "each rule carries a fixed severity and category, and FR-16, which states that each rule carries a "
          "message template. One row per rule closes both. The severity column of this table is the answer to "
          "\"where does a finding's severity come from?\".")
h3("C.1 Rule-engine rules (source = rule)")
table(["Rule id", "Category", "Severity", "Base", "Message template"],
      [("hardcoded-secret", "security", "Critical", "8",
        "A credential-like value is assigned to {symbol} - move it to an environment variable and rotate the key."),
       ("sql-concat", "security", "High", "5",
        "SQL is built by string concatenation in {symbol}() - use a parameterised query."),
       ("dangerous-eval", "security", "High", "5",
        "{symbol}() calls {construct} on runtime input - replace it with an explicit parser or dispatch table."),
       ("complex-function", "code-design", "Medium", "3",
        "{symbol}() has cyclomatic complexity {value}, over the limit of {threshold} - split it into smaller functions."),
       ("long-method", "code-design", "Medium", "3",
        "{symbol}() is {value} lines long, over the limit of {threshold} - extract cohesive blocks into helpers."),
       ("deep-nesting", "code-design", "Medium", "3",
        "{symbol}() nests {value} levels deep, over the limit of {threshold} - use early returns to flatten it."),
       ("duplicate-block", "code-design", "Low", "1",
        "This block is duplicated {value} times across the file - extract it into a shared helper."),
       ("large-file", "code-design", "Low", "1",
        "{file} is {value} lines long, over the limit of {threshold} - consider splitting it by responsibility.")],
      [1700, 1300, 900, 600, 4860], "Table 4-2. Rule register: rule-engine rules.")

h3("C.2 SATD marker patterns (source = satd)")
body_text("The category is predicted by ML-1; the severity comes from the marker matched in the comment text "
          "(FR-9.2). Patterns are evaluated high to low and the highest match wins.")
table(["Marker pattern", "Severity", "Base", "Message template"],
      [("\\b(FIXME|BUG|XXX|BROKEN|DO\\s*NOT\\s*(SHIP|MERGE))\\b", "High", "5",
        "Self-admitted defect: '{comment_text}' - classified as {predicted_category}."),
       ("\\b(TODO|HACK|TEMP|TEMPORARY|WORKAROUND|KLUDGE|REFACTOR)\\b", "Medium", "3",
        "Self-admitted debt: '{comment_text}' - classified as {predicted_category}."),
       ("\\b(NOTE|REVIEW|NIT|IDEA|QUESTION|MAYBE)\\b", "Low", "1",
        "Self-admitted note: '{comment_text}' - classified as {predicted_category}."),
       ("(no marker matched - the classifier detected debt in prose alone)", "Medium", "3",
        "Self-admitted debt: '{comment_text}' - classified as {predicted_category}.")],
      [3400, 900, 600, 4460], "Table 4-3. Rule register: SATD marker patterns.")

h3("C.3 Debt-category taxonomy and dataset label mapping")
body_text("Normative for FR-9.3. The five predicted categories are fixed by the labels present in "
          "satd-dataset-code_comments.csv - the only dataset file v1.0 uses (FR-9.1). The mapping is applied once, "
          "in the machine-learning service's post-processing; the normalised product value is what is stored and "
          "displayed. This table closes decision D5.")
table(["Category (product value)", "Dataset label", "Instances", "Assigned by"],
      [("code-design", "code/design_debt", "2,703", "ML-1 and the rule engine"),
       ("requirement", "requirement_debt", "757", "ML-1"),
       ("**defect**", "**defect_debt**", "**472**", "ML-1"),
       ("test", "test_debt", "85", "ML-1"),
       ("documentation", "documentation_debt", "54", "ML-1"),
       ("security", "(not present in the dataset)", "-", "Rule engine only - never predicted"),
       ("(not a category)", "non_debt", "58,204", "The negative class of the debt / not-debt decision")],
      [2400, 2600, 1200, 3160], "Table 4-4. Debt-category taxonomy and dataset label mapping.")
body_text("defect debt is a developer admitting a known bug - the dataset's own example is "
          "\"// FIXME formatters are not thread-safe\". It is distinct from the risk model: ML-2 predicts future "
          "bug-proneness from numeric metrics, whereas defect debt is a current, admitted defect stated in prose. "
          "It is included because it is the third-largest debt class in the comment data, larger than test and "
          "documentation combined.")
info("[Note for the reader: only 6.54% of the 62,275 labelled comments are debt at all, and documentation (0.09%) "
     "and test (0.14%) are severely under-represented. See FR-25 for the reporting obligation this creates.]")

h3("C.4 File-level risk message")
body_text("Not a finding - this is the badge and tooltip text for the per-file risk score (FR-10).")
table(["Trigger", "Template"],
      [("risk_score shown on a file",
        "High-risk file ({risk}): {salient_signals}. For example: \"High-risk file (0.78): high complexity (CCN 18) "
        "and frequent change (14 commits/90d).\"")],
      [2600, 6760], "Table 4-5. Risk badge template.")
info("[TO FILL - expand C.1 as further rules are added (target approximately 30 to 50 rows as language coverage "
     "grows); confirm the {construct} and {salient_signals} interpolation fields with the backend.]")

h2("Appendix D - Referencing conventions")
bullet("Refer to all data and information in a standard format (IEEE referencing style).", )
bullet("For algorithms, techniques and theories, refer to textbooks and peer-reviewed articles.")
bullet("For tools, refer to the official web page and include \"(Accessed on <date>)\".")
bullet("For similar work, refer to research articles that describe the work.")
doc.add_paragraph(style="Body Text")

doc.save(OUT)

# ───────────────────────── 7. make Word refresh the TOC on open ──────────────
tmp = OUT + ".tmp"
with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if re.match(r"word/(header|footer)\d+\.xml$", item.filename):
            # placeholders inside w:fldSimple are not reachable through python-docx runs
            s = data.decode("utf-8")
            for k, v in REPL.items():
                s = s.replace(k.replace("<", "&lt;").replace(">", "&gt;"), v)
            data = s.encode("utf-8")
        if item.filename == "word/settings.xml":
            s = data.decode("utf-8")
            if "updateFields" not in s:
                for anchor in ("<w:hdrShapeDefaults", "<w:footnotePr", "<w:endnotePr",
                               "<w:compat", "<w:rsids", "</w:settings>"):
                    if anchor in s:
                        s = s.replace(anchor, '<w:updateFields w:val="true"/>' + anchor, 1)
                        break
            data = s.encode("utf-8")
        zout.writestr(item, data)
os.replace(tmp, OUT)

d2 = docx.Document(OUT)
print("OK ->", OUT)
print("paragraphs:", len(d2.paragraphs), " tables:", len(d2.tables), " sections:", len(d2.sections))
print("size KB:", round(os.path.getsize(OUT) / 1024, 1))
