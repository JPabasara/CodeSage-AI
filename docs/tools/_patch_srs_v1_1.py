# -*- coding: utf-8 -*-
"""Produce SRS v1.1 by patching the submitted v1.0 file in place.

    python docs/tools/_patch_srs_v1_1.py

Input  : docs/Deliverables/SRS/v1.0/Software_Requirement_Specification_v1.0.docx
Output : docs/Deliverables/SRS/v1.1/Software_Requirement_Specification_v1.1.docx

Everything not listed below is inherited byte-for-byte, including styles,
headers, footers, the table of contents field and the list of tables.

What v1.1 changes, and why:

1. Identity moves to Asgardeo. v1.0 says the browser signs in with GitHub OAuth
   and carries a bearer token per request. Both are now wrong. The application is
   an OIDC client of Asgardeo, GitHub is federated *inside* Asgardeo, and the
   browser holds an httpOnly session cookie rather than any token.
2. FastAPI is the Backend-for-Frontend. The OIDC exchange happens in the backend
   and provider tokens never leave it. Four new security requirements state that
   as a requirement instead of leaving it to implementation.
3. The four auth rows of Table 3.106 were rotated - each path carried the *next*
   row's purpose - so the endpoint table did not describe an implementable API.
4. A read-latency requirement is added. FR-21 derives every score on read, which
   makes dashboard-read latency the load-bearing performance risk, and DBR-32
   already referred to a target that did not exist.
"""
from __future__ import annotations

import os
import shutil
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402

from _docx_patch import (  # noqa: E402
    clone_row,
    find_paragraph,
    find_table,
    insert_paragraph_after,
    insert_rows,
    must_replace,
    patch_header_text,
    row_index,
    set_row,
    verify_docx,
)

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
SRC = os.path.join(ROOT, "docs", "Deliverables", "SRS", "v1.0",
                   "Software_Requirement_Specification_v1.0.docx")
OUT_DIR = os.path.join(ROOT, "docs", "Deliverables", "SRS", "v1.1")
OUT = os.path.join(OUT_DIR, "Software_Requirement_Specification_v1.1.docx")

os.makedirs(OUT_DIR, exist_ok=True)

# This script REBUILDS v1.1 from v1.0. Once v1.1 has been opened in Word - fields
# refreshed, figures replaced - re-running it would silently throw that work away.
# Refuse instead. A later version is a NEW patch script reading v1.1, never a
# re-run of this one.
if os.path.exists(OUT):
    raise SystemExit(
        f"{OUT} already exists.\n"
        "This script rebuilds v1.1 from v1.0 and would overwrite any Word edits "
        "(refreshed fields, replaced figures) already made to it.\n"
        "Delete the file first if you really mean to rebuild it from scratch."
    )

shutil.copyfile(SRC, OUT)
doc = docx.Document(OUT)
changes: list[str] = []


def note(label: str, hits: int) -> None:
    changes.append(f"{label} ({hits})")


# ══════════════════════════════════════════════════════════════════════════════
# 1. Title page and document properties
# ══════════════════════════════════════════════════════════════════════════════
note("title page version", must_replace(doc, "Version 1.0", "Version 1.1", whole=True))
note("title page date", must_replace(doc, "Date: 9th August 2026",
                                     "Date: 15th August 2026", whole=True))

doc.core_properties.title = "Software Requirements Specification"
doc.core_properties.subject = "Code Sage AI"
doc.core_properties.author = "Group 16"
doc.core_properties.category = "CS3203 Software Engineering Project"
doc.core_properties.revision = 11

# ══════════════════════════════════════════════════════════════════════════════
# 2. Revision History - keep the v1.0 row, add v1.1
# ══════════════════════════════════════════════════════════════════════════════
rev = doc.tables[0]
clone_row(rev, len(rev.rows) - 1)
set_row(rev, len(rev.rows) - 1, [
    "15/Aug/2026",
    "1.1",
    "User authentication moved from direct GitHub OAuth to Asgardeo, with GitHub "
    "federated inside Asgardeo, so that additional sign-in methods can be added "
    "without further application changes (FR-1). The backend was designated the "
    "Backend-for-Frontend: it completes the sign-in exchange and keeps identity "
    "tokens server-side, and the browser holds only a session cookie. Four "
    "security requirements were added for that boundary (SEC-17 to SEC-20), and a "
    "dashboard read-latency requirement was added (PERF-11). The application "
    "endpoint table was corrected: the four authentication rows had been listed "
    "against the wrong descriptions, the paths are now provider-neutral, and the "
    "scan-history endpoint records its snapshot parameter. The REST contract is "
    "now named as a repository artefact rather than described in prose.",
    "Group 16",
])
changes.append("revision history row 1.1")

# ══════════════════════════════════════════════════════════════════════════════
# 3. Section 1.3 - glossary terms for the new identity model
# ══════════════════════════════════════════════════════════════════════════════
glossary = find_table(doc, "Technical debt (TD)", "SATD")
insert_rows(glossary, len(glossary.rows) - 1, [
    ["Asgardeo",
     "The hosted identity provider used for sign-in. GitHub, and later other "
     "methods, are federated inside Asgardeo, so the application integrates with "
     "one provider rather than one per sign-in method (FR-1)."],
    ["OIDC",
     "OpenID Connect, the identity layer over OAuth 2.0 that Asgardeo and the "
     "backend use to establish who a user is."],
    ["PKCE",
     "Proof Key for Code Exchange (RFC 7636). Binds an authorization code to the "
     "request that started the sign-in, so an intercepted code cannot be redeemed "
     "by anyone else."],
    ["BFF",
     "Backend-for-Frontend. The pattern in which the browser talks only to this "
     "system's own backend, and that backend holds every external credential. In "
     "v1.0 the FastAPI service is the BFF (SEC-17)."],
    ["Session",
     "A server-side record of one signed-in user. The browser holds only an "
     "opaque identifier for it, in an httpOnly cookie, so the session can be "
     "revoked at sign-out (SEC-10)."],
])
changes.append("glossary: 5 identity terms")

# ══════════════════════════════════════════════════════════════════════════════
# 4. Section 1.4 - references
# ══════════════════════════════════════════════════════════════════════════════
ref13 = find_paragraph(doc, "Replication Package of Deep Learning and Data Augmentation")
insert_paragraph_after(
    ref13,
    '[16] Internet Engineering Task Force, "Proof Key for Code Exchange by OAuth '
    'Public Clients," RFC 7636, Sep. 2015. [Online]. Available: '
    'https://www.rfc-editor.org/rfc/rfc7636. [Accessed: Aug. 14, 2026].',
)
insert_paragraph_after(
    ref13,
    '[15] OpenID Foundation, "OpenID Connect Core 1.0." [Online]. Available: '
    'https://openid.net/specs/openid-connect-core-1_0.html. '
    '[Accessed: Aug. 14, 2026].',
)
insert_paragraph_after(
    ref13,
    '[14] WSO2 LLC, "Asgardeo Documentation." [Online]. Available: '
    'https://wso2.com/asgardeo/docs/. [Accessed: Aug. 14, 2026].',
)
changes.append("references [14][15][16]")

# ══════════════════════════════════════════════════════════════════════════════
# 5. Section 2.4 - assumptions no longer require a GitHub account
# ══════════════════════════════════════════════════════════════════════════════
note("assumption: sign-in", must_replace(
    doc,
    "Users have valid GitHub accounts and can authenticate through the supported "
    "GitHub authentication mechanism.",
    "Users can authenticate through the identity provider using one of the "
    "sign-in methods it offers. In v1.0 that is a GitHub account, federated "
    "inside Asgardeo; adding another method is a configuration change in the "
    "identity provider and not a change to this system.",
    whole=True))

note("assumption: provider availability", must_replace(
    doc,
    "GitHub services required for authentication, repository metadata retrieval "
    "and repository cloning remain available.",
    "The identity provider remains available for sign-in, and the GitHub services "
    "required for repository metadata retrieval and repository cloning remain "
    "available. An outage of either is an external-service failure under REL-01.",
    whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 6. FR-1 - the functional requirement itself
# ══════════════════════════════════════════════════════════════════════════════
fr = find_table(doc, "FR-1", "FR-25")
note("FR-1", must_replace(
    doc,
    "Sign in with GitHub; authenticated session; land on Projects; sign out from "
    "the Account menu. Session state is managed per request (token), never held "
    "in server memory.",
    "Sign in through Asgardeo, which federates GitHub; land on Projects; sign out "
    "from the Account menu. The backend performs the sign-in exchange, records a "
    "server-side session and returns an httpOnly session cookie to the browser. "
    "Identity-provider tokens are never sent to the browser (SEC-17). Signing out "
    "deletes the server-side session, so the cookie stops being accepted "
    "immediately (SEC-10). On first sign-in the system creates the user's "
    "workspace and its default Balanced scoring profile.",
    whole=True))

note("FR-2 identity key", must_replace(
    doc,
    "Every project scan and finding belongs to exactly one workspace.",
    "Every project, scan and finding belongs to exactly one workspace. A user is "
    "identified by the stable subject identifier issued by the identity provider, "
    "never by an email address, because an email address can change while that "
    "identifier cannot."))

# ══════════════════════════════════════════════════════════════════════════════
# 7. Section 3.4 - the missing read-latency requirement
# ══════════════════════════════════════════════════════════════════════════════
perf = find_table(doc, "PERF-01", "PERF-10")
insert_rows(perf, len(perf.rows) - 1, [
    ["PERF-11", "Dashboard Read Latency",
     "Because every score is derived when results are read rather than stored "
     "(FR-21), the dashboard payload for one branch, including the derivation of "
     "its trend history, shall be returned within 2 seconds at the 95th "
     "percentile under the baseline capacity defined in PERF-06."],
])
changes.append("PERF-11 read latency")

note("DBR-32 pointer", must_replace(
    doc,
    "within the response-time requirements defined in the Performance "
    "Requirements section",
    "within the response-time requirements defined in the Performance "
    "Requirements section, and PERF-11 in particular"))

# ══════════════════════════════════════════════════════════════════════════════
# 8. Section 3.5 - security. SEC-10 amended, SEC-17 to SEC-20 added
# ══════════════════════════════════════════════════════════════════════════════
sec = find_table(doc, "SEC-01", "SEC-16")

note("SEC-10", must_replace(
    doc,
    "Authenticated sessions shall be invalidated when the user logs out or when "
    "the associated authorization is revoked. Session credentials shall not "
    "remain valid beyond their configured expiration period.",
    "Sessions shall be held server-side, so that signing out deletes the session "
    "record and the browser's cookie stops being accepted on the very next "
    "request. A self-contained token that stays valid until it expires shall not "
    "be used as the session credential, because it cannot be revoked. Sessions "
    "shall also expire after a configured idle period and a configured absolute "
    "lifetime, whichever comes first.",
    whole=True))

insert_rows(sec, len(sec.rows) - 1, [
    ["SEC-17", "Authentication Boundary",
     "The authorization-code exchange with the identity provider shall be "
     "performed by the backend, using PKCE. Identity, access and refresh tokens "
     "shall be held server-side only, and shall never be transmitted to the "
     "browser, written to browser storage or placed in a URL. The browser shall "
     "receive only an opaque session identifier in an httpOnly, Secure, "
     "SameSite=Lax cookie."],
    ["SEC-18", "Deny by Default",
     "Every application endpoint shall require an authenticated session. The only "
     "exceptions shall be the two endpoints that begin and complete sign-in and "
     "the liveness probe, and that exemption list shall be stated explicitly in "
     "the REST contract. An endpoint added without an authorization decision "
     "shall fail closed rather than be reachable anonymously."],
    ["SEC-19", "Request Forgery and Origin Control",
     "The sign-in flow shall carry a single-use, signed state value that is "
     "verified on return. Cross-origin access shall be restricted to an explicit "
     "list of permitted origins; wildcard origins shall not be used where "
     "credentials are allowed. State-changing requests shall be rejected when "
     "their origin is not on that list."],
    ["SEC-20", "Security Response Headers",
     "Responses shall carry HTTP Strict-Transport-Security, a Content-Security-"
     "Policy, X-Content-Type-Options, a policy forbidding the application from "
     "being framed by another site, and a Referrer-Policy that does not leak "
     "internal paths to third parties."],
])
changes.append("SEC-17..SEC-20")

# ══════════════════════════════════════════════════════════════════════════════
# 9. Section 3.10.3 - Table 3.104, external interfaces
# ══════════════════════════════════════════════════════════════════════════════
ext = find_table(doc, "GitHub REST API", "Git (clone)")
set_row(ext, row_index(ext, "GitHub OAuth"), [
    "Asgardeo (identity provider)",
    "User authentication (FR-1). GitHub is federated inside Asgardeo, so this "
    "system integrates with one provider rather than one per sign-in method.",
    "Outbound; HTTPS, OpenID Connect authorization-code flow with PKCE, completed "
    "by the backend (SEC-17).",
])
note("Table 3.104 GitHub REST purpose", must_replace(
    doc,
    "Repository metadata: name, owner, visibility, default branch, branch list "
    "and head commit SHAs.",
    "Repository metadata: name, owner, visibility, default branch, branch list "
    "and head commit SHAs. Not used for authentication.",
    whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 10. Table 3.105 - name the contract artefact
# ══════════════════════════════════════════════════════════════════════════════
note("Table 3.105 REST row", must_replace(
    doc,
    "HTTPS with JSON; payload formats are governed by the data contract; "
    "specified as an OpenAPI document in the Software Architecture / Design "
    "Document.",
    "HTTPS with JSON, in snake_case. The normative contract is the OpenAPI 3.1 "
    "document held in the repository at docs/api/openapi.yaml; the frontend's "
    "TypeScript types and the backend's request and response models are both "
    "generated from or checked against it. The browser calls this API directly: "
    "there is no separate proxy tier in front of it.",
    whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 11. Table 3.106 - the endpoint table
# ══════════════════════════════════════════════════════════════════════════════
api = find_table(doc, "/api/profiles/active", "/api/projects")

set_row(api, row_index(api, "/api/auth/github", "Begin sign-in"), [
    "GET", "/api/auth/login",
    "Begin sign-in. Redirects the browser to the identity provider, carrying a "
    "signed single-use state value and a PKCE challenge (FR-1, SEC-19).",
])
set_row(api, row_index(api, "/api/auth/github/login"), [
    "GET", "/api/auth/callback",
    "The identity provider's redirect target. Verifies state, exchanges the "
    "authorization code in the backend, creates the server-side session, sets the "
    "session cookie and redirects the browser to the Projects page (FR-1, SEC-17).",
])
set_row(api, row_index(api, "/api/auth/github/callback"), [
    "GET", "/api/auth/session",
    "Return the signed-in user, or 401 when there is no valid session (FR-1).",
])
set_row(api, row_index(api, "/api/auth/session", "End the session"), [
    "POST", "/api/auth/logout",
    "Delete the server-side session and clear the cookie (FR-1, SEC-10).",
])
changes.append("Table 3.106: four auth rows corrected")

set_row(api, row_index(api, "/api/repos/{repoId}/scans"), [
    "GET", "/api/repos/{repo_id}/scans?branch=",
    "Scan history: the stored snapshot summaries for one branch (FR-19). Each "
    "entry carries the snapshot identifier that GET .../health accepts as "
    "?snapshot_id= to load a past scan into the dashboard.",
])
note("health endpoint snapshot_id", must_replace(
    doc,
    "Full dashboard payload for one branch snapshot (FR-12 to FR-18). Scores are "
    "derived under the workspace's active profile.",
    "Full dashboard payload for one branch snapshot (FR-12 to FR-18). Scores are "
    "derived under the workspace's active profile, which the server resolves "
    "itself; no profile is passed as a parameter. An optional ?snapshot_id= loads "
    "a past snapshot instead of the latest one (FR-19).",
    whole=True))

# Path parameters are snake_case on the wire, so the table should spell them the
# same way the contract does.
for old, new in (("{repoId}", "{repo_id}"), ("{scanId}", "{scan_id}")):
    for row in api.rows:
        cell = row.cells[1]
        if old in cell.text:
            from _docx_patch import set_paragraph_text
            set_paragraph_text(cell.paragraphs[0], cell.text.replace(old, new))
changes.append("Table 3.106: snake_case path parameters")

# ══════════════════════════════════════════════════════════════════════════════
# 12. Table 3.107 - communications
# ══════════════════════════════════════════════════════════════════════════════
comms = find_table(doc, "Scan progress", "Backend to Git host")
set_row(comms, row_index(comms, "Session and credentials"), [
    "Session and credentials",
    "httpOnly, Secure, SameSite=Lax session cookie",
    "The cookie carries an opaque identifier for a server-side session, never a "
    "token. JavaScript cannot read it, so a cross-site scripting fault cannot "
    "steal it, and sign-out revokes it server-side (SEC-10, SEC-17).",
])
insert_rows(comms, len(comms.rows) - 1, [
    ["Backend to identity provider",
     "HTTPS; OpenID Connect authorization-code flow with PKCE",
     "Outbound from the backend only. The browser never contacts the token "
     "endpoint (SEC-17)."],
])
changes.append("Table 3.107: session + identity provider rows")

# ══════════════════════════════════════════════════════════════════════════════
# 13. Database requirements touched by the identity change
# ══════════════════════════════════════════════════════════════════════════════
dbr = find_table(doc, "DBR-1", "DBR-35")
note("DBR-4", must_replace(
    doc,
    "The database shall store workspaces, users and the membership relationship "
    "between them. In v1.0, each workspace shall be limited to one active member. "
    "However, the data model shall support multiple memberships per workspace in "
    "a later release.",
    "The database shall store workspaces, users and the membership relationship "
    "between them. A user record shall be keyed on the stable subject identifier "
    "issued by the identity provider; the sign-in method used and any display "
    "details are stored as attributes of that user, not as its identity. In v1.0 "
    "each workspace shall be limited to one active member, and the data model "
    "shall support multiple memberships per workspace in a later release.",
    whole=True))

insert_rows(dbr, len(dbr.rows) - 1, [
    ["DBR-36", "Session Storage",
     "The database shall store one record per authenticated session, holding its "
     "identifier, the user and workspace it belongs to, its creation time, its "
     "last-used time and its expiry. Deleting the record shall be sufficient to "
     "end the session (SEC-10). No identity-provider token shall be stored in "
     "plaintext (DBR-29, SEC-08)."],
])
changes.append("DBR-36 session storage")

# ══════════════════════════════════════════════════════════════════════════════
# 14. Applicable standards
# ══════════════════════════════════════════════════════════════════════════════
std = find_table(doc, "ST-1", "ST-6")
note("ST-2", must_replace(
    doc,
    "OAuth 2.0 / RFC 6749: Governs the standards for secure user authentication, "
    "authorization, and data extraction.",
    "OpenID Connect Core 1.0, over OAuth 2.0 (RFC 6749) with PKCE (RFC 7636): "
    "governs how the system authenticates a user through the identity provider "
    "and how the authorization code is protected in transit."))

# ══════════════════════════════════════════════════════════════════════════════
# 15. Appendix B - name the artefact the shapes live in
# ══════════════════════════════════════════════════════════════════════════════
note("Appendix B", must_replace(
    doc,
    "are defined by the shared data contract that lives in the OpenAPI contract, "
    "which is the single source for their field names, types and permitted values. "
    "Any change to those shapes is a change to this appendix.",
    "are defined by the OpenAPI contract held in the repository at "
    "docs/api/openapi.yaml, which is the single normative source for their field "
    "names, types and permitted values. This appendix is normative in that it "
    "names that contract and maps each output to the requirement it serves; the "
    "shapes themselves are defined there."))

# ══════════════════════════════════════════════════════════════════════════════
# 16. Corrections carried over from the v1.0 review
# ══════════════════════════════════════════════════════════════════════════════
note("REL-13 wording", must_replace(
    doc, "it is classified as a critical debt", "it is classified as a critical defect"))

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)

# ══════════════════════════════════════════════════════════════════════════════
# 17. Header: version string and document identifier live in header XML
# ══════════════════════════════════════════════════════════════════════════════
note("header version + doc id", patch_header_text(OUT, [
    ("CS3203-G16-SRS-v1.0", "CS3203-G16-SRS-v1.1"),
    ("1.0", "1.1"),
    ("09/08/2026", "15/08/2026"),
]))

# The file is only finished once it is something Word will actually open.
verify_docx(OUT)

print(f"SRS v1.1 written to {OUT}")
for line in changes:
    print("  -", line)
