# -*- coding: utf-8 -*-
"""Shared helpers for producing v1.1 of a deliverable **by patching v1.0 in place**.

Why patch rather than regenerate. `_build_srs_docx.py` and `_build_sad_docx.py`
build a document from the course template. They were correct for v1.0, but the
team then edited the result by hand for days — figures were embedded, tables were
rewritten, captions were renumbered. Re-running a generator would throw all of
that away.

So v1.1 is produced by copying the v1.0 file and editing it. Every style,
header, footer, margin, numbering definition, embedded image and field code is
inherited untouched, because it is literally the same file.

The one thing that needs care is run splitting. Word stores a paragraph as a list
of runs, and it splits them wherever formatting, spell-check state or an edit
session changed - so the visible sentence "Sign in with GitHub" may live in four
runs. A naive run-by-run replace silently misses those. `set_paragraph_text`
below rewrites the whole paragraph, keeping the first run so its formatting
survives.
"""
from __future__ import annotations

import copy
import re
import zipfile

from docx.oxml.ns import qn


# ── reading ──────────────────────────────────────────────────────────────────

def iter_paragraphs(doc):
    """Every paragraph in the document, including those nested inside tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def norm(text: str) -> str:
    """Collapse whitespace and normalise the punctuation Word substitutes.

    Word turns quotes into curly quotes and hyphens into en-dashes as you type,
    so a match string copied out of the .md will not equal the .docx text unless
    both are folded first.
    """
    text = (text.replace("’", "'").replace("‘", "'")
                .replace("“", '"').replace("”", '"')
                .replace("–", "-").replace("—", "-")
                .replace(" ", " "))
    return re.sub(r"\s+", " ", text).strip()


# ── writing ──────────────────────────────────────────────────────────────────

def set_paragraph_text(paragraph, text: str) -> None:
    """Replace a paragraph's text, keeping its style and its first run's format.

    Field codes are never touched: a paragraph holding a `fldChar` is skipped by
    the callers below, because rewriting one would destroy the TOC.
    """
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra._r.getparent().remove(extra._r)
    else:
        paragraph.add_run(text)


def has_field(paragraph) -> bool:
    """True if the paragraph contains a Word field (TOC entry, PAGEREF, SEQ)."""
    xml = paragraph._p.xml
    return "fldChar" in xml or "instrText" in xml


def replace_text(doc, old: str, new: str, *, count: int = 0, whole: bool = False) -> int:
    """Substitute `old` with `new` across the body and every table cell.

    `whole=True` matches the entire paragraph rather than a substring, which is
    what you want for a table cell you are rewriting wholesale.

    Returns the number of paragraphs changed, so a caller can assert that an
    edit actually landed instead of silently doing nothing.
    """
    hits = 0
    target = norm(old)
    for paragraph in iter_paragraphs(doc):
        if has_field(paragraph):
            continue
        current = norm(paragraph.text)
        if whole:
            if current != target:
                continue
            set_paragraph_text(paragraph, new)
        else:
            if target not in current:
                continue
            set_paragraph_text(paragraph, current.replace(target, new))
        hits += 1
        if count and hits >= count:
            break
    return hits


def must_replace(doc, old: str, new: str, *, count: int = 0, whole: bool = False,
                 label: str = "") -> int:
    """`replace_text`, but a miss is an error rather than a silent no-op."""
    hits = replace_text(doc, old, new, count=count, whole=whole)
    if hits == 0:
        raise AssertionError(f"no match for {label or old[:70]!r}")
    return hits


# ── tables ───────────────────────────────────────────────────────────────────

def find_table(doc, *needles, skip: int = 0):
    """The first table whose text contains every needle (after `skip` matches)."""
    wanted = [norm(n) for n in needles]
    seen = 0
    for table in doc.tables:
        blob = norm(" ".join(c.text for r in table.rows for c in r.cells))
        if all(w in blob for w in wanted):
            if seen == skip:
                return table
            seen += 1
    raise AssertionError(f"no table matching {needles!r}")


def row_index(table, *needles) -> int:
    wanted = [norm(n) for n in needles]
    for i, row in enumerate(table.rows):
        blob = norm(" ".join(c.text for c in row.cells))
        if all(w in blob for w in wanted):
            return i
    raise AssertionError(f"no row matching {needles!r} in table")


def set_row(table, index: int, values: list[str]) -> None:
    """Overwrite a row's cells, leaving the row's own formatting alone.

    A cell may legitimately hold several paragraphs; the first receives the text
    and the rest are emptied rather than deleted, so cell padding is unchanged.
    """
    cells = table.rows[index].cells
    for cell, value in zip(cells, values):
        paragraphs = cell.paragraphs
        set_paragraph_text(paragraphs[0], value)
        for extra in paragraphs[1:]:
            set_paragraph_text(extra, "")


def clone_row(table, template_index: int, at_index: int | None = None):
    """Copy a row (so it inherits its borders and shading) and insert it.

    `at_index=None` appends. Otherwise the copy is inserted *before* that row.
    """
    source = table.rows[template_index]._tr
    new = copy.deepcopy(source)
    if at_index is None:
        table._tbl.append(new)
    else:
        table.rows[at_index]._tr.addprevious(new)
    return new


def insert_rows(table, template_index: int, rows: list[list[str]],
                at_index: int | None = None) -> None:
    """Insert several rows at once, each cloned from the same template row."""
    for offset, values in enumerate(rows):
        position = None if at_index is None else at_index + offset
        clone_row(table, template_index, position)
        target = len(table.rows) - 1 if at_index is None else at_index + offset
        set_row(table, target, values)


def delete_row(table, index: int) -> None:
    tr = table.rows[index]._tr
    tr.getparent().remove(tr)


# ── paragraphs ───────────────────────────────────────────────────────────────

def find_paragraph(doc, needle: str, *, style: str | None = None, skip: int = 0):
    target = norm(needle)
    seen = 0
    for paragraph in doc.paragraphs:
        if style and paragraph.style.name != style:
            continue
        if target in norm(paragraph.text):
            if seen == skip:
                return paragraph
            seen += 1
    raise AssertionError(f"no paragraph containing {needle[:70]!r}")


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    """Add a paragraph directly below `paragraph`, in the same document."""
    new_p = copy.deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:fldSimple")):
            new_p.remove(child)
    paragraph._p.addnext(new_p)

    from docx.text.paragraph import Paragraph
    created = Paragraph(new_p, paragraph._parent)
    if style:
        created.style = paragraph.part.document.styles[style]
    created.add_run(text)
    return created


def insert_paragraphs_after(paragraph, blocks: list[tuple[str, str | None]]):
    """Insert (text, style) pairs in order below `paragraph`. Returns the last."""
    anchor = paragraph
    for text, style in blocks:
        anchor = insert_paragraph_after(anchor, text, style)
    return anchor


def insert_paragraphs_before(paragraph, blocks: list[tuple[str, str | None]]):
    """Insert (text, style) pairs in order *above* `paragraph`.

    Needed when a new subsection has to land at the end of the section it belongs
    to: the last block of that section is often a table, and a table is not a
    paragraph, so the only reliable anchor is the heading that follows it.
    """
    from docx.text.paragraph import Paragraph

    first = insert_paragraph_after(paragraph, blocks[0][0], blocks[0][1])
    paragraph._p.addprevious(first._p)
    anchor = Paragraph(first._p, paragraph._parent)
    for text, style in blocks[1:]:
        anchor = insert_paragraph_after(anchor, text, style)
    return anchor


def patch_header_text(path: str, subs: list[tuple[str, str]],
                      pattern: str = r"word/header\d*\.xml") -> int:
    """Substitute visible text inside header/footer parts, on the raw XML.

    python-docx cannot do this safely: merely *reading* `section.header` creates an
    empty definition and breaks inheritance from the template. So the part is
    rewritten in the zip instead.

    ⚠️ The substitution is applied **only inside `<w:t>` elements**, never to the
    whole file. A blanket string replace over header XML looks harmless and is not:
    replacing "1.0" with "1.1" to bump a version number also rewrites the XML
    declaration to `<?xml version="1.1"?>`, and Word refuses to open a document
    whose parts claim XML 1.1. lxml accepts it, so the damage survives every check
    short of opening the file in Word.

    Returns the number of parts changed.
    """
    with zipfile.ZipFile(path) as archive:
        items = {name: archive.read(name) for name in archive.namelist()}

    changed = 0
    for name in items:
        if not re.match(pattern, name):
            continue
        xml = items[name].decode("utf8")

        def substitute(match: "re.Match[str]") -> str:
            text = match.group(2)
            for old, new in subs:
                text = text.replace(old, new)
            return f"{match.group(1)}{text}{match.group(3)}"

        patched = re.sub(r"(<w:t(?:\s[^>]*)?>)(.*?)(</w:t>)", substitute, xml,
                         flags=re.DOTALL)
        if patched != xml:
            items[name] = patched.encode("utf8")
            changed += 1

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, blob in items.items():
            archive.writestr(name, blob)
    return changed


def verify_docx(path: str) -> None:
    """Fail loudly if the file is not something Word will open.

    Written after a bad header patch produced a document that passed every other
    check - the zip CRCs were fine and python-docx opened it happily - and still
    could not be opened in Word. Each assertion below corresponds to a way that
    happened.
    """
    from lxml import etree

    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad is not None:
            raise AssertionError(f"corrupt zip entry: {bad}")

        names = archive.namelist()
        for required in ("[Content_Types].xml", "word/document.xml"):
            if required not in names:
                raise AssertionError(f"missing required part: {required}")

        for name in names:
            if not name.endswith(".xml") and not name.endswith(".rels"):
                continue
            blob = archive.read(name)

            # Word parses XML 1.0 only. lxml will happily accept 1.1, so this has
            # to be checked on the raw bytes rather than after parsing.
            head = blob[:200].decode("utf8", "ignore")
            if "<?xml" in head and 'version="1.0"' not in head \
                    and "version='1.0'" not in head:
                declaration = head.split("?>")[0] + "?>"
                raise AssertionError(
                    f"{name}: XML declaration is not version 1.0 - Word will "
                    f"refuse to open this file.\n    {declaration}"
                )

            try:
                etree.fromstring(blob)
            except etree.XMLSyntaxError as error:
                raise AssertionError(f"{name}: malformed XML - {error}") from error

        # Every part must have a declared content type, by extension or by name.
        # python-docx rewrites `Default` extension rules into per-part `Override`
        # rules when it saves, which is equivalent - unless one gets dropped, and
        # then Word reports the same unhelpful "error trying to open the file".
        types_ns = "{http://schemas.openxmlformats.org/package/2006/content-types}"
        root = etree.fromstring(archive.read("[Content_Types].xml"))
        defaults = {e.get("Extension").lower()
                    for e in root if e.tag == f"{types_ns}Default"}
        overrides = {e.get("PartName").lstrip("/")
                     for e in root if e.tag == f"{types_ns}Override"}

        uncovered = [
            name for name in names
            if name != "[Content_Types].xml" and not name.endswith("/")
            and name.rsplit(".", 1)[-1].lower() not in defaults
            and name not in overrides
        ]
        if uncovered:
            raise AssertionError(
                "parts with no declared content type (Word will refuse the file): "
                + ", ".join(uncovered)
            )


def revise(doc, rows: list[tuple[str, str, str, str]]) -> None:
    """Rewrite the Revision History table so it carries the new version row.

    The table is located by content, not by index, because both deliverables put
    an "Individual Contributions" table immediately after it and the two are easy
    to confuse.
    """
    table = find_table(doc, "Revision History") if _has_header(doc, "Revision History") \
        else doc.tables[0]
    while len(table.rows) - 1 < len(rows):
        clone_row(table, len(table.rows) - 1)
    for i, values in enumerate(rows, start=1):
        set_row(table, i, list(values))


def _has_header(doc, text: str) -> bool:
    target = norm(text)
    return any(target in norm(t.rows[0].cells[0].text) for t in doc.tables if t.rows)
