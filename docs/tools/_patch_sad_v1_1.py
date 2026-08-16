# -*- coding: utf-8 -*-
"""Produce SAD v1.1 by patching the submitted v1.0 file in place.

    python docs/tools/_patch_sad_v1_1.py

Input  : docs/Deliverables/SAD/v1.0/Software_Architecture_Document_v1.0.docx
Output : docs/Deliverables/SAD/v1.1/Software_Architecture_Document_v1.1.docx

Every embedded figure, style, header, footer and field code is inherited
byte-for-byte, because this is the same file with edits applied. The figures
themselves are NOT redrawn here - the diagram corrections are listed in
docs/NEXT_STEPS.md and are done by hand in draw.io.

What v1.1 changes:

1. Identity. Sign-in is Asgardeo (OIDC, GitHub federated), and the FastAPI
   service is the Backend-for-Frontend: it completes the exchange, holds the
   provider tokens and hands the browser only a session cookie. A new Section 6.4
   describes that flow the way 6.2 describes the profile write.
2. Arithmetic. Five category weights plus one trust slider is six numbers. Three
   places said seven, left over from when there were six categories.
3. Table names. The prose referred to SCAN and FILE_SCORE, which are not in the
   data model - it holds ANALYSIS_ATTEMPT and SNAPSHOT, and per-file facts live
   in STATIC_METRIC, PROCESS_METRIC and BUG_RISK_PREDICTION.
4. The glossary said "exactly five values" and then listed six, including the
   `defect` category that was removed with the SATDAUG corpus change.
5. Cross-references to SRS Table 3.96 and 3.97 are now 3.106 and 3.107.
6. The architectural style is named in Section 2. A reader could previously
   finish the document without learning whether this is a monolith or a set of
   services.
7. Three table cells were blank, and Section 1.5 described sections 10 to 12
   incorrectly.
"""
from __future__ import annotations

import os
import shutil
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402

from _docx_patch import (  # noqa: E402
    clone_row,
    find_paragraph,
    find_table,
    insert_paragraph_after,
    insert_paragraphs_after,
    insert_paragraphs_before,
    insert_rows,
    must_replace,
    patch_header_text,
    replace_text,
    row_index,
    set_row,
    verify_docx,
)

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
SRC = os.path.join(ROOT, "docs", "Deliverables", "SAD", "v1.0",
                   "Software_Architecture_Document_v1.0.docx")
OUT_DIR = os.path.join(ROOT, "docs", "Deliverables", "SAD", "v1.1")
OUT = os.path.join(OUT_DIR, "Software_Architecture_Document_v1.1.docx")

os.makedirs(OUT_DIR, exist_ok=True)

# See the note in _patch_srs_v1_1.py: this rebuilds v1.1 from v1.0, so re-running
# it after the figures have been replaced in Word would destroy that work.
if os.path.exists(OUT):
    raise SystemExit(
        f"{OUT} already exists.\n"
        "This script rebuilds v1.1 from v1.0 and would overwrite any Word edits "
        "(refreshed fields, replaced figures) already made to it.\n"
        "Delete the file first if you really mean to rebuild it from scratch."
    )

shutil.copyfile(SRC, OUT)
doc = docx.Document(OUT)
changes: list[str] = []


def note(label: str, hits: int) -> None:
    changes.append(f"{label} ({hits})")


BODY = "Body Text" if "Body Text" in [s.name for s in doc.styles] else None

# ══════════════════════════════════════════════════════════════════════════════
# 1. Title page
# ══════════════════════════════════════════════════════════════════════════════
note("title version", must_replace(doc, "Version 1.0", "Version 1.1", whole=True))
note("title date", must_replace(doc, "Date: 9th August 2026",
                                "Date: 15th August 2026", whole=True))

doc.core_properties.title = "Software Architecture Document"
doc.core_properties.subject = "Code Sage AI"
doc.core_properties.author = "Group 16"
doc.core_properties.category = "CS3203 Software Engineering Project"
doc.core_properties.revision = 11

# ══════════════════════════════════════════════════════════════════════════════
# 2. Revision History
# ══════════════════════════════════════════════════════════════════════════════
rev = doc.tables[0]
clone_row(rev, len(rev.rows) - 1)
set_row(rev, len(rev.rows) - 1, [
    "15/Aug/2026",
    "1.1",
    "Aligned with SRS v1.1. Sign-in becomes Asgardeo with GitHub federated, and "
    "the API service is named as the Backend-for-Frontend that holds the identity "
    "tokens; a new Section 6.4 sets out that exchange. Section 2 now names the "
    "architectural style and states why a set of independent services was not "
    "chosen. The scoring profile is corrected to six numbers, being five category "
    "weights and one trust slider. Prose that referred to SCAN and FILE_SCORE "
    "tables now uses the entities the data model actually defines. Security and "
    "performance tables gain the requirements added in SRS v1.1, and three empty "
    "table cells and the Section 1.5 section list were corrected.",
    "Group 16",
])
changes.append("revision history row 1.1")

# ══════════════════════════════════════════════════════════════════════════════
# 3. Section 1.2 Scope
# ══════════════════════════════════════════════════════════════════════════════
note("scope paragraph", must_replace(
    doc,
    "This document covers the architecture of the v1.0 release: a multi-tenant "
    "web application in which a user signs in with GitHub, connects a public "
    "repository by pasting its URL,",
    "This document covers the architecture of the v1.0 release: a multi-tenant "
    "web application in which a user signs in through the Asgardeo identity "
    "provider, connects a public repository by pasting its URL,"))

note("scope limit 1", must_replace(
    doc,
    "Repository access in v1.0 is public repositories by URL only. Sign-in uses "
    "GitHub OAuth, and the analysis pipeline reads the repository through an "
    "ordinary git clone.",
    "Repository access in v1.0 is public repositories by URL only. Sign-in uses "
    "Asgardeo, which federates GitHub, and the analysis pipeline reads the "
    "repository through an ordinary git clone. Signing in and reading a "
    "repository are therefore two unrelated grants: the system never asks a user "
    "for repository permission, and never holds a GitHub token on their behalf."))

# ══════════════════════════════════════════════════════════════════════════════
# 4. Section 1.3 glossary
# ══════════════════════════════════════════════════════════════════════════════
note("glossary category", must_replace(
    doc,
    "What kind of debt a finding is. Exactly five values: code-design, "
    "requirement, defect, documentation, test and security (SRS FR-9.3).",
    "What kind of debt a finding is. Exactly five values: code-design, "
    "requirement, documentation, test and security (SRS FR-9.3).",
    whole=True))

glossary = find_table(doc, "Finding", "Visibility floor")
insert_rows(glossary, len(glossary.rows) - 1, [
    ["Asgardeo",
     "The hosted identity provider. GitHub sign-in is federated inside it, so "
     "this system is a client of one provider rather than of each sign-in method "
     "(SRS FR-1)."],
    ["BFF",
     "Backend-for-Frontend. The browser talks only to this system's own API, and "
     "that API holds every external credential. The FastAPI service plays that "
     "role here (SRS SEC-17)."],
    ["Session",
     "A server-side record of one signed-in user. The browser holds only an "
     "opaque identifier for it, in an httpOnly cookie, so sign-out can revoke it "
     "(SRS SEC-10)."],
])
changes.append("glossary: Asgardeo, BFF, Session")

# ══════════════════════════════════════════════════════════════════════════════
# 5. Table 1-2 references - name the contract artefact
# ══════════════════════════════════════════════════════════════════════════════
refs = find_table(doc, "Software Requirements Specification", "Change Request CR-001")
note("reference 6", must_replace(
    doc,
    "Shared data contract - the single source of truth for the shapes crossing "
    "the frontend, backend and database boundary (SRS SP-4).",
    "REST contract - the OpenAPI 3.1 document that is the single source of truth "
    "for every shape crossing the frontend and backend boundary (SRS SP-4). The "
    "frontend's TypeScript types are generated from it.",
    whole=True))
note("reference 6 path", must_replace(
    doc, "apps/web/src/lib/types/index.ts", "docs/api/openapi.yaml", whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 6. Section 1.5 Overview - the section list was wrong
# ══════════════════════════════════════════════════════════════════════════════
note("overview section list", must_replace(
    doc,
    "Section 10 covers sizing and performance targets, Section 11 explains how "
    "the architecture delivers the quality attributes, and Section 12 lists the "
    "references.",
    "Section 10 covers sizing and performance targets, Section 11 records the "
    "measurements to be taken during performance testing, and Section 12 explains "
    "how the architecture delivers each quality attribute."))

note("overview figure count", must_replace(
    doc,
    "Ten figures are referenced by number throughout. Each one is cited in the "
    "text as \"Figure n\" and is described in the sentences that follow it.",
    "Figures are referenced by number throughout. Figures 1 to 10 carry the "
    "architecture itself, and Figures 1.1 to 1.7 expand the individual use cases "
    "of Figure 1. Each is cited in the text as \"Figure n\" and is described in "
    "the sentences that follow it.",
    whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 7. Section 2 - name the architectural style
# ══════════════════════════════════════════════════════════════════════════════
note("typo: sperate", must_replace(doc, "A sperate Data View", "A separate Data View"))

anchor = find_paragraph(doc, "It follows the 4+1 architectural view approach")
insert_paragraphs_after(anchor, [
    ("The architectural style is a modular monolith with an asynchronous worker "
     "and one extracted inference service. The API and the worker are built from "
     "one codebase and run from the same container image with different commands, "
     "so a change that spans a request handler and the pipeline it queues is one "
     "release rather than two. Inside that monolith four further styles are used, "
     "each for a specific requirement: the code is layered, so the presentation "
     "layer cannot reach the database; the scan pipeline is a pipe and filter "
     "arrangement of clone, extract, detect and finalize, so the cancellation "
     "check has well-defined places to sit; the API and its workers form a "
     "producer and competing-consumer pair over a broker, so concurrency is a "
     "deployment setting rather than a code change; and the write path is "
     "separated from the read path, so a scan stores facts and every score is "
     "derived when the dashboard is requested.", BODY),
    ("A set of independently deployed services was considered and rejected, for "
     "four reasons that hold regardless of team size. The domain is one bounded "
     "context: repository, branch, analysis attempt, snapshot and finding form a "
     "single joined graph that the dashboard reads in one query. The central "
     "correctness rule is a database transaction, because SRS DBR-22 requires a "
     "snapshot and all of its findings, metrics and predictions to be committed "
     "together or not at all; split across services that becomes a distributed "
     "transaction with compensating actions, and the half-written snapshot that "
     "SRS FR-6 forbids becomes a state the system can actually reach. Tenant "
     "isolation depends on PostgreSQL Row-Level Security, which is one mechanism "
     "over one database and becomes several separate arguments once the data is "
     "split. And the boundaries that would be drawn here follow workload rather "
     "than business capability, which is a reason to separate processes, not to "
     "separate services and their data.", BODY),
    ("The machine-learning service is the one process deployed separately, and it "
     "is not a service in that sense either: it owns no data and no business "
     "capability. It is extracted so that its unavailability is a designed mode "
     "rather than an exception, so that a retrained model can be replaced by "
     "swapping a mounted artefact, and so that the training dependencies stay out "
     "of the API image. Section 7 sets out the resulting deployment.", BODY),
])
changes.append("Section 2: architectural style named")

# ══════════════════════════════════════════════════════════════════════════════
# 8. Table 3-1 - a goal for the authentication boundary
# ══════════════════════════════════════════════════════════════════════════════
goals = find_table(doc, "G1", "G7")
insert_rows(goals, len(goals.rows) - 1, [
    ["G8",
     "No credential the browser holds may be worth stealing.",
     "Sign-in is completed by the API, not by the browser. The API is the "
     "Backend-for-Frontend: it performs the authorization-code exchange with "
     "Asgardeo, keeps the identity tokens in its own process and database, and "
     "returns an httpOnly session cookie that JavaScript cannot read. Sessions "
     "are server-side, so signing out revokes access on the next request (SRS "
     "SEC-10, SEC-17 to SEC-20)."],
])
changes.append("G8 authentication boundary goal")

note("constraints stack", must_replace(
    doc,
    "PostgreSQL for storage; Docker for deployment.",
    "PostgreSQL for storage; Asgardeo for identity; Docker for deployment."))

# ══════════════════════════════════════════════════════════════════════════════
# 9. Section 4 - the sign-in use case
# ══════════════════════════════════════════════════════════════════════════════
note("UC-1 name", must_replace(
    doc, "Sign In / GitHub Authentication", "Sign In (Asgardeo, GitHub federated)",
    whole=True))
note("UC-1 significance", must_replace(
    doc,
    "Establishes authenticated access to the system and provides the security "
    "foundation for repository access and role-based features.",
    "Establishes authenticated access and fixes the trust boundary for the whole "
    "system: the exchange with the identity provider happens in the API, so no "
    "external credential ever reaches the browser.",
    whole=True))

note("4.1 narrative", must_replace(
    doc,
    "The Sign In process begins when a developer initiates authentication and the "
    "system sends the user to the GitHub Authentication Service, where the "
    "credentials are checked for verification by GitHub. Once they are verified, "
    "the system obtains an authentication token and confirms credentials and "
    "create user session If there is a failure in the authentication process or "
    "if the authentication token has failed, then the system will display "
    "authentication error.",
    "The Sign In process begins when a developer asks to sign in. The API "
    "redirects the browser to Asgardeo, which presents the sign-in methods "
    "configured for the workspace; in v1.0 that is GitHub, federated inside "
    "Asgardeo, so GitHub verifies the credentials and Asgardeo, not this system, "
    "receives the result. Asgardeo then returns the browser to the API's callback "
    "with an authorization code. The API verifies the state value it issued at "
    "the start, exchanges the code for the identity token over its own "
    "server-to-server connection, creates a session row and sets an httpOnly "
    "session cookie on the browser. On a first sign-in it also creates the user's "
    "workspace and its default scoring profile, so no later read has to cope with "
    "a workspace that has none. If the state does not verify, or the exchange "
    "fails, no session is created, an audit record is written and the user is "
    "returned to the sign-in page with an error.",
    whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 10. Section 5 - classes and packages
# ══════════════════════════════════════════════════════════════════════════════
packages = find_table(doc, "Identity and Workspace Management")
note("identity package classes", must_replace(
    doc, "User, Workspace, Membership, AuthenticationService",
    "User, Workspace, Membership, Session, AuthenticationService, "
    "IdentityProviderGateway", whole=True))
note("identity package role", must_replace(
    doc,
    "Manages authenticated users, workspace boundaries, and user-to-workspace "
    "membership. It establishes the tenant context used to isolate repositories, "
    "analyses, and other workspace-owned data.",
    "Manages authenticated users, workspace boundaries, and user-to-workspace "
    "membership. It completes the sign-in exchange with the identity provider, "
    "holds the resulting session server-side, and establishes the tenant context "
    "used to isolate repositories, analyses, and other workspace-owned data.",
    whole=True))

# The role cell for Dashboard and Reporting was left empty in v1.0.
set_row(packages, row_index(packages, "Dashboard and Reporting"), [
    "Dashboard and Reporting",
    "DashboardService",
    "Assembles stored analysis facts and dynamically derived scores into the "
    "read-only representations the dashboard needs, including current and "
    "historical repository health and technical-debt information.",
])
changes.append("filled empty cell: Dashboard and Reporting")

# Needle chosen to be unique to Table 5-2: the design-package table above also
# names AuthenticationService and ScoringEngine, so those would match it first.
classes = find_table(doc, "Represents an authenticated user of the platform")

note("AuthenticationService", must_replace(
    doc, "Handles GitHub authentication, session validation, and sign-out behavior.",
    "Completes the OpenID Connect exchange with Asgardeo, creates and validates "
    "sessions, and revokes them at sign-out.", whole=True))
note("AuthenticationService note", must_replace(
    doc, "Stateless application service.",
    "The application service is stateless, but the session it issues is not: the "
    "session record lives in the database so that signing out ends it "
    "immediately (SRS SEC-10).", whole=True))
note("GitHubGateway note", must_replace(
    doc, "Handles OAuth, repository metadata retrieval, branch information, and "
         "repository cloning.",
    "Handles repository metadata retrieval, branch information and repository "
    "cloning. It plays no part in authentication: identity comes from Asgardeo.",
    whole=True))

# The responsibility cell for Finding was left empty in v1.0.
set_row(classes, row_index(classes, "Finding", "source rule or satd"), [
    "Finding",
    "Represents one atomic technical-debt issue detected in source code, at a "
    "file, line and symbol.",
    "A finding has source rule or satd; category and severity are fixed at "
    "detection time.",
])
changes.append("filled empty cell: Finding responsibility")

insert_rows(classes, len(classes.rows) - 1, [
    ["Session",
     "Represents one signed-in user's session, held server-side.",
     "The browser receives only an opaque identifier for it. Deleting the record "
     "ends the session on the next request."],
    ["IdentityProviderGateway",
     "Encapsulates every call to the identity provider.",
     "Builds the authorization request, verifies the returned state and exchanges "
     "the authorization code. It is the only component that sees a provider "
     "token."],
])
changes.append("Session + IdentityProviderGateway classes")

note("backend subsystem BFF", must_replace(
    doc,
    "FastAPI handles synchronous user-facing operations, while Celery workers "
    "perform long-running repository analyses asynchronously.",
    "FastAPI handles synchronous user-facing operations and is the only part of "
    "the system the browser talks to, while Celery workers perform long-running "
    "repository analyses asynchronously."))

# ══════════════════════════════════════════════════════════════════════════════
# 11. Section 6 - entity names the data model actually defines
# ══════════════════════════════════════════════════════════════════════════════
note("6: last successful scan", must_replace(
    doc,
    "A cancelled or failed scan leaves a SCAN row that has no FINDING and no "
    "FILE_SCORE rows, because the worker stopped before the persist stage. If the "
    "head SHA were compared against such a row, the system would skip the work "
    "and serve a snapshot that was never written. The same qualifier applies to "
    "the dashboard read, which resolves the latest snapshot where phase = done.",
    "A cancelled or failed scan leaves an ANALYSIS_ATTEMPT row and no SNAPSHOT at "
    "all, because the worker stopped before the finalize stage. That is why the "
    "comparison is made against the SHA of the most recent attempt that produced "
    "a snapshot: comparing against any attempt would let the system skip the work "
    "and then serve a snapshot that was never written. The split between the two "
    "tables is what makes this structural rather than a condition every query has "
    "to remember."))

note("6: API answers first", must_replace(
    doc,
    "POST /api/repos/{repoId}/scan inserts the SCAN row, enqueues the job and "
    "returns a scan identifier with phase = queued. The client then polls GET "
    "/api/repos/{repoId}/scan/{scanId} once per second (SRS Table 3.97).",
    "POST /api/repos/{repo_id}/scan inserts the ANALYSIS_ATTEMPT row, enqueues "
    "the job and returns a scan identifier with phase = queued. The client then "
    "polls GET /api/repos/{repo_id}/scan/{scan_id} once per second (SRS Table "
    "3.107)."))

note("6: phase in postgres", replace_text(
    doc, "The status endpoint reads SCAN.phase from the database",
    "The status endpoint reads ANALYSIS_ATTEMPT.phase from the database"))
note("6: terminal phase", replace_text(
    doc, "Every terminal phase is therefore written to SCAN by the process that "
         "reaches it.",
    "Every terminal phase is therefore written to ANALYSIS_ATTEMPT by the process "
    "that reaches it."))
note("6: cancel writes phase", replace_text(
    doc, "deletes its clone and writes phase = cancelled",
    "deletes its clone and writes phase = cancelled to its ANALYSIS_ATTEMPT row"))
note("6: cancel poll row", replace_text(
    doc, "because the worker writes phase = cancelled to the same row the status "
         "endpoint already reads",
    "because the worker writes phase = cancelled to the same attempt row the "
    "status endpoint already reads"))
note("6: failed scan", must_replace(
    doc,
    "The worker writes phase = error and the error message onto the existing SCAN "
    "row. This satisfies SRS SP-13: a failure reported by a user can be diagnosed "
    "from the database without reading server logs. Nothing was written to "
    "FINDING or FILE_SCORE, so the previous snapshot is unaffected and remains "
    "what the dashboard shows.",
    "The worker writes phase = error and the error message onto the existing "
    "ANALYSIS_ATTEMPT row. This satisfies SRS SP-13: a failure reported by a user "
    "can be diagnosed from the database without reading server logs. No SNAPSHOT "
    "was created, so the previous snapshot is unaffected and remains what the "
    "dashboard shows."))

note("6: figure 5 insert", must_replace(
    doc,
    "Only when the two differ does it insert a SCAN row and enqueue a job.",
    "Only when the two differ does it insert an ANALYSIS_ATTEMPT row and enqueue "
    "a job."))
note("6: figure 5 worker writes", must_replace(
    doc,
    "It records its phase by writing to SCAN and its progress by publishing to "
    "Redis.",
    "It records its phase by writing to ANALYSIS_ATTEMPT and its progress by "
    "publishing to Redis."))

note("6: ML fallback wording", replace_text(
    doc,
    "no SATD findings appear, and every risk factor falls back to 1.0 so that no "
    "finding receives ML-2 risk scores.",
    "no SATD findings appear, and no risk score is recorded for any file. A "
    "missing risk score is reported as absent rather than as zero, because zero "
    "would mean the file was assessed and found safe; with no score the risk "
    "factor falls back to 1.0 and boosts nothing."))

# ══════════════════════════════════════════════════════════════════════════════
# 12. Section 6.2 - six numbers, not seven
# ══════════════════════════════════════════════════════════════════════════════
note("6.2 write transaction", must_replace(
    doc,
    "It updates weights and trust_s on SCORE_PROFILE, and sets active_profile_id "
    "on WORKSPACE. Seven numbers change. That is the whole write.",
    "It writes the five weights and trust_s onto the workspace's SCORING_PROFILE "
    "row and marks that row active. Six numbers change. That is the whole write. "
    "Exactly one profile can be active per workspace because a partial unique "
    "index on the workspace column, restricted to active rows, makes a second one "
    "impossible to insert - a guarantee the database holds rather than one the "
    "application has to remember."))

note("6.2 figure decision", must_replace(
    doc,
    "The write is a single transaction covering seven numbers. The five category "
    "weights and the trust slider are written to SCORE_PROFILE at the same moment "
    "that active_profile_id is set on WORKSPACE, so that no reader can ever "
    "observe one without the other. No SCAN, FINDING or FILE_SCORE row is touched "
    "anywhere in the exchange, which is precisely what SRS FR-21 means when it "
    "states that changing a profile creates no snapshot.",
    "The write is a single transaction covering six numbers. The five category "
    "weights and the trust slider are written to SCORING_PROFILE and that row is "
    "marked active in the same transaction, so no reader can ever observe one "
    "without the other. No ANALYSIS_ATTEMPT, SNAPSHOT or FINDING row is touched "
    "anywhere in the exchange, which is precisely what SRS FR-21 means when it "
    "states that changing a profile creates no snapshot."))

note("6.2 SRS table refs", replace_text(doc, "SRS Table 3.96", "SRS Table 3.106"))

profile_paths = find_table(doc, "Run a scan", "Apply a profile")
set_row(profile_paths, row_index(profile_paths, "Writes"), [
    "Writes",
    "A new ANALYSIS_ATTEMPT and, on success, an immutable SNAPSHOT with its "
    "FINDING, metric and prediction rows",
    "Six numbers on one SCORING_PROFILE row",
])
set_row(profile_paths, row_index(profile_paths, "POST /api/repos"), [
    "Request",
    "POST /api/repos/{repo_id}/scan",
    "PUT /api/profiles/active",
])
changes.append("Table 6-3 corrected")

note("6.2 health url", replace_text(
    doc, "GET /api/repos/{repoId}/health?branch", "GET /api/repos/{repo_id}/health?branch"))

# ══════════════════════════════════════════════════════════════════════════════
# 13. New Section 6.4 - signing in
# ══════════════════════════════════════════════════════════════════════════════
# Section 6.3 ends with Table 6-4, and a table is not a paragraph, so the only
# stable anchor for "end of Section 6" is the heading that starts Section 7.
deployment_heading = find_paragraph(doc, "Deployment View", style="Heading 1")

insert_paragraphs_before(deployment_heading, [
    ("Signing in", "Heading 2"),
    ("Sign-in is the third exchange worth describing on its own, because it is "
     "where the system's trust boundary is drawn. The requirement is SRS FR-1, "
     "and SRS SEC-17 to SEC-20 constrain how it may be met.", BODY),
    ("The exchange runs in the API, not in the browser. When a user asks to sign "
     "in, the API generates a single-use state value and a PKCE verifier, keeps "
     "both server-side, and redirects the browser to Asgardeo carrying the state "
     "and the derived challenge. Asgardeo presents whichever sign-in methods the "
     "tenant has configured - GitHub in v1.0 - and returns the browser to the "
     "API's callback with an authorization code. The API checks that the state "
     "matches the one it issued, then exchanges the code and the verifier for an "
     "identity token over its own connection to Asgardeo. The browser is not part "
     "of that exchange and never sees the token.", BODY),
    ("What the browser receives instead is a session cookie. The API creates a "
     "session row, and sets a cookie holding nothing but that row's identifier, "
     "marked httpOnly so JavaScript cannot read it, Secure so it travels only "
     "over TLS, and SameSite=Lax so another site cannot cause the browser to send "
     "it on a state-changing request. Three properties follow. A cross-site "
     "scripting fault cannot steal a credential the page cannot read. Signing out "
     "deletes the session row, so the cookie stops being accepted on the very "
     "next request, which is what SRS SEC-10 requires and what a self-contained "
     "token could not offer without a revocation list that would be most of a "
     "session store anyway. And because no provider token is ever stored in the "
     "browser, there is nothing in the browser worth stealing.", BODY),
    ("Authorization is decided before the handler runs. Every route except the "
     "two that begin and complete sign-in, and the liveness probe, resolves the "
     "session first and rejects the request with 401 when there is none. The same "
     "dependency then binds the caller's workspace to the database transaction, "
     "so Row-Level Security has a tenant to filter on. Because that binding is "
     "the same dependency that opens the session, a handler cannot accidentally "
     "run without a tenant: it would have no database session to use. This is "
     "what makes SRS SEC-18 structural rather than a convention.", BODY),
    ("What the system stores about a user is deliberately small: the stable "
     "subject identifier issued by Asgardeo, a display name and an email address. "
     "There is no password, because the system never sees one, and no GitHub "
     "token, because v1.0 clones public repositories anonymously and therefore "
     "never acts at GitHub on a user's behalf. Keying the user on the subject "
     "identifier rather than on the email address matters for the same reason: an "
     "email address can be changed by its owner, and the subject identifier "
     "cannot.", BODY),
    ("Two consequences reach the rest of the architecture. Adding a sign-in "
     "method later - Google, or a username and password for the reviewers and "
     "stakeholders of a later release - is a configuration change inside "
     "Asgardeo, because this system integrates with the provider and not with "
     "each method behind it. And the identity provider is a hosted dependency, so "
     "an outage of it stops new sign-ins while existing sessions continue to "
     "work; SRS REL-01 excludes external-service outages from the availability "
     "target for exactly this class of dependency.", BODY),
])
changes.append("Section 6.4 Signing in")

# ══════════════════════════════════════════════════════════════════════════════
# 14. Section 7 - deployment
# ══════════════════════════════════════════════════════════════════════════════
nodes = find_table(doc, "Client device", "Worker container(s)")
note("backend node connections", must_replace(
    doc,
    "Communicates with the frontend through HTTPS/JSON REST; Redis for task "
    "enqueueing/progress; PostgreSQL for persistent data, GitHub through "
    "HTTPS/OAuth/REST as required.",
    "Receives HTTPS/JSON REST requests from the browser and is the only backend "
    "process reachable from outside the private network. Uses Redis for task "
    "enqueueing and progress, PostgreSQL for persistent data, Asgardeo over HTTPS "
    "for the sign-in exchange, and the GitHub REST API for repository metadata.",
    whole=True))
note("backend node processes", must_replace(
    doc,
    "FastAPI application. Handles authentication, repository/project operations,",
    "FastAPI application, acting as the Backend-for-Frontend. Handles "
    "authentication and session issue, repository/project operations,"))

insert_rows(nodes, len(nodes.rows) - 1, [
    ["Identity provider (external)",
     "Asgardeo, hosting the sign-in methods and issuing identity tokens. Not "
     "deployed by this project.",
     "Reached over HTTPS by the backend container only. The browser is redirected "
     "to it during sign-in but never exchanges a token with it directly."],
])
changes.append("deployment: identity provider node")

# ══════════════════════════════════════════════════════════════════════════════
# 15. Section 9 - data view
# ══════════════════════════════════════════════════════════════════════════════
groups = find_table(doc, "Tenant and access data")
note("group 1 entities", must_replace(
    doc, "User, Membership, Workspace, SecurityAuditRecord",
    "User, Membership, Workspace, Session, SecurityAuditRecord", whole=True))

entities = find_table(doc, "Workspace", "AnalysisAttempt", "ScoringProfile")

# The "what it holds" cell for Branch was left empty in v1.0.
set_row(entities, row_index(entities, "Branch"), [
    "Branch",
    "Repository branches, including branch name, current head commit SHA, and "
    "whether the branch is the default branch.",
    "Stored fact",
])
changes.append("filled empty cell: Branch")

note("membership entity", must_replace(
    doc,
    "GitHub-linked user information and the relationship between users and "
    "workspaces.",
    "The relationship between users and workspaces. A user is identified by the "
    "subject identifier issued by the identity provider; the sign-in method used "
    "is an attribute, not the key."))

insert_rows(entities, len(entities.rows) - 1, [
    ["Session",
     "One record per signed-in user: the session identifier the cookie carries, "
     "the user and workspace it belongs to, and its creation, last-use and expiry "
     "times. No identity-provider token is stored in plaintext.",
     "Stored fact"],
])
changes.append("Session entity")

rules = find_table(doc, "FINDING.severity", "SCORE_PROFILE.weights")
set_row(rules, row_index(rules, "SCAN (append-only)"), [
    "SNAPSHOT (append-only)",
    "Finalized analysis results are immutable and append-only. A successful "
    "analysis creates a new snapshot rather than updating an earlier one, so "
    "trend, history and delta are queries over existing rows in chronological "
    "order. A failed or cancelled attempt keeps its ANALYSIS_ATTEMPT row for "
    "diagnosis but creates no snapshot, which is what makes \"only a completed "
    "scan can be read back\" a property of the schema rather than a condition "
    "every query must remember.",
    "FR-19, FR-21, REL-05, DBR-22, DBR-23",
])
set_row(rules, row_index(rules, "WORKSPACE.active_profile_id"), [
    "SCORING_PROFILE.is_active",
    "Marks the profile currently used when deriving priorities, health scores, "
    "grades, trends and category breakdowns for the workspace. A partial unique "
    "index on the workspace column, restricted to rows where this flag is set, "
    "makes a second active profile impossible to insert. Changing which profile "
    "is active changes derived values only; stored analysis facts are untouched.",
    "FR-20, FR-21",
])
note("SCORE_PROFILE naming", replace_text(
    doc, "SCORE_PROFILE.weights and trust_s", "SCORING_PROFILE weights and trust_s"))
changes.append("Table 9-2: SNAPSHOT and is_active rules")

note("scoring profile entity", must_replace(
    doc,
    "Workspace-configurable category weights and rule-versus-model trust setting "
    "used when deriving priorities and health scores.",
    "Workspace-configurable weights, one per debt category, and the "
    "rule-versus-model trust setting used when deriving priorities and health "
    "scores. Five weights and one trust value, six numbers in all."))

# ══════════════════════════════════════════════════════════════════════════════
# 16. Section 10 - the read-latency target
# ══════════════════════════════════════════════════════════════════════════════
dims = find_table(doc, "Registered users at baseline", "Scan enqueue")
insert_rows(dims, len(dims.rows) - 1, [
    ["Dashboard read latency",
     "Return the dashboard payload for one branch, including the derivation of "
     "its trend history, within 2 seconds at the 95th percentile under the "
     "baseline capacity.",
     "PERF-11"],
])
changes.append("Table 10-1: PERF-11")

note("10 efficient retrieval", must_replace(
    doc,
    "Therefore, loading the dashboard or changing the active scoring profile does "
    "not require the repository to be scanned again.",
    "Therefore, loading the dashboard or changing the active scoring profile does "
    "not require the repository to be scanned again. Because that derivation runs "
    "on every read, SRS PERF-11 sets a budget for it, and Section 11 records the "
    "measurement that confirms it."))

# ══════════════════════════════════════════════════════════════════════════════
# 17. Section 12 - quality attributes
# ══════════════════════════════════════════════════════════════════════════════
perfq = find_table(doc, "PERF-01", "Provide immediate feedback")
insert_rows(perfq, len(perfq.rows) - 1, [
    ["PERF-11", "Keep the derived dashboard fast",
     "One indexed query per dashboard read followed by an in-memory scoring pass, "
     "with per-group sums available as an optimisation should the measurement in "
     "Section 11 show it is needed."],
])
changes.append("Table 12-2: PERF-11")

sec = find_table(doc, "Secure user authentication")
note("SEC-01 mechanism", must_replace(
    doc, "Users authenticate through GitHub OAuth before accessing system features.",
    "Users authenticate through Asgardeo, which federates GitHub. The API "
    "completes the exchange and issues a server-side session; the browser holds "
    "only an httpOnly cookie.", whole=True))

insert_rows(sec, len(sec.rows) - 1, [
    ["SEC-10", "Revoke access at sign-out",
     "Sessions are rows in the database, so signing out deletes the row and the "
     "cookie stops being accepted on the next request. A self-contained token "
     "would stay valid until it expired."],
    ["SEC-17", "Keep identity tokens away from the browser",
     "The API is the Backend-for-Frontend: it performs the authorization-code "
     "exchange with PKCE and keeps every provider token in its own process. The "
     "browser receives an opaque session identifier and nothing else."],
    ["SEC-18", "Authenticate every endpoint by default",
     "Session resolution is a dependency of the same object that opens the "
     "database transaction, so a handler cannot run without an authenticated "
     "caller and a bound tenant. Only sign-in start, sign-in callback and the "
     "liveness probe are exempt, and the contract lists them."],
    ["SEC-19", "Resist request forgery and unwanted origins",
     "A signed single-use state value is verified on return from the identity "
     "provider; the session cookie is SameSite=Lax; and cross-origin access is "
     "limited to an explicit list of origins rather than a wildcard."],
    ["SEC-20", "Harden the responses themselves",
     "Strict-Transport-Security, a Content-Security-Policy, X-Content-Type-"
     "Options, a frame-ancestors restriction and a Referrer-Policy are set on "
     "responses by middleware, so the policy is applied in one place."],
])
changes.append("Table 12-3: SEC-10, SEC-17..20")

note("PRI-02", must_replace(
    doc, "GitHub tokens and API credentials are stored only within backend services.",
    "No user credential is stored at all: sign-in happens at the identity "
    "provider and v1.0 clones public repositories anonymously. Service "
    "credentials are held only by backend containers, supplied as environment "
    "variables.", whole=True))

# ══════════════════════════════════════════════════════════════════════════════
# 18. Remove empty headings left behind by editing
# ══════════════════════════════════════════════════════════════════════════════
# An empty Heading 2 between two use-case subsections was producing a blank entry
# in the table of contents.
removed = 0
for para in list(doc.paragraphs):
    if para.style.name.startswith("Heading") and not para.text.strip():
        para._p.getparent().remove(para._p)
        removed += 1
note("empty headings removed", removed)

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)


# ══════════════════════════════════════════════════════════════════════════════
# 18. Header: version and document identifier
# ══════════════════════════════════════════════════════════════════════════════
note("header version + doc id", patch_header_text(OUT, [
    ("CS3203-G16-SAD-v1.0", "CS3203-G16-SAD-v1.1"),
    ("1.0", "1.1"),
    ("09/08/2026", "15/08/2026"),
]))

# The file is only finished once it is something Word will actually open.
verify_docx(OUT)

print(f"SAD v1.1 written to {OUT}")
for line in changes:
    print("  -", line)
