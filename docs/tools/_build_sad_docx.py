# -*- coding: utf-8 -*-
"""
⚠️ HISTORICAL. This script built the v0.x SKELETON from the blank course template.
It is superseded by the patch approach in _patch_srs_v1_1.py / _patch_sad_v1_1.py,
which edit the previous version instead of rebuilding from the template - because
rebuilding throws away every figure and hand-written table added since.

Kept because it is the record of how the document structure was first laid out,
and because it is the only thing that knows how to inherit the template's styles,
headers, footers and numbering from scratch. See docs/tools/README.md.

Build the Code Sage AI SAD .docx ON TOP OF the course template (#4) so that every
style, header, footer, margin and numbering definition is literally the
template's own (styles.xml / numbering.xml / header*.xml / footer*.xml / sectPr
are all inherited untouched).

This produces a TEMPLATE, not a finished document:
  - Section 1 (Introduction) is fully written, because it states the project's
    scope decisions and those must not be re-guessed by whoever fills the rest.
  - Sections 2-12 are headings + a short plain-English prompt + empty grids +
    numbered figure captions. No RUP boilerplate.
  - Figures are numbered 1-10 with captions, but no figure is generated. The
    numbering matches Software_Architecture_Document.md exactly.

Blue bracketed text (style=InfoBlue) is guidance and must be deleted before
submission, per the template's own convention.
"""
import re, shutil, zipfile, os
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ───────────────────────────────────────────────────────────────────
# Repo-relative, so the script runs on any machine and from any directory.
# The course templates stayed in docs/Templates/ when the scripts moved here.
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
TPL = os.path.join(ROOT, "docs", "Templates", "4 = Template for Software Architecture Document.docx")
OUT = os.path.join(ROOT, "docs", "Deliverables", "_generated", "Software_Architecture_Document_skeleton.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

# Writing into _generated/ rather than over a deliverable is deliberate: this
# script produces a SKELETON from the blank course template, not a finished
# document. Pointing it at the real filename would make one careless run replace
# a document the team had spent days editing.
if os.path.exists(OUT):
    raise SystemExit(
        f"{OUT} already exists - delete it first if you really want to rebuild "
        "the skeleton from the blank course template."
    )



TBLPR = (
    '<w:tblPr %s>'
    '<w:tblW w:w="0" w:type="auto"/>'
    '<w:tblBorders>'
    '<w:top w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:left w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:right w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="single" w:sz="6" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
    '<w:tblLayout w:type="fixed"/>'
    '<w:tblLook w:val="0000" w:firstRow="0" w:lastRow="0" w:firstColumn="0" '
    'w:lastColumn="0" w:noHBand="0" w:noVBand="0"/>'
    '</w:tblPr>' % nsdecls('w')
)

shutil.copyfile(TPL, OUT)
doc = docx.Document(OUT)
body = doc.element.body

# ───────────────────────── 1. placeholder replacement ────────────────────────
REPL = {
    "<Project Name>": "Code Sage AI",
    "<Company Name>": "Group 16",
    "<Subsystem or Feature>": "the AI-Powered Technical-Debt Analytics Dashboard",
    "<1.0>": "1.0",
    "<dd/mmm/yy>": "09/Aug/2026",
    "<document identifier>": "CS3203-G16-SAD-v1.0",
}


def fix_runs(container):
    for p in container.paragraphs:
        for r in p.runs:
            for k, v in REPL.items():
                if k in r.text:
                    r.text = r.text.replace(k, v)
    for t in container.tables:
        for row in t.rows:
            for c in row.cells:
                fix_runs(c)


fix_runs(doc)
# Headers/footers are patched at the XML level at the bottom of this script.
# Touching section.first_page_header through python-docx CREATES an empty
# definition and breaks inheritance from the template.

doc.core_properties.title = "Software Architecture Document"
doc.core_properties.subject = "Code Sage AI"
doc.core_properties.author = "Group 16"
doc.core_properties.category = "CS3203 Software Engineering Project"

kids = list(body.iterchildren())

# ───────────────────────── 2. title page ─────────────────────────────────────
# blocks 6 and 7 are the RUP "how to use this template" notes -> remove
for i in (6, 7):
    body.remove(kids[i])

# group identification, inserted after "Version 1.0"
anchor = kids[3]
for line, sz, bold in [
    ("Group 16   |   Project ID 7   |   CS3203 - Software Engineering Project", 12, True),
    ("Mentor: Mr. Anju Chamantha", 11, False),
    ("Date: 9 August 2026", 11, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(sz)
    r.bold = bold
    anchor.addnext(p._p)
    anchor = p._p

# ───────────────────────── 3. Revision History table ─────────────────────────
rev = doc.tables[0]
rows = [
    ("22/Jul/2026", "0.1",
     "Initial architecture: 4+1 views, data model and quality attributes for v1.0.", "Group 16"),
    ("30/Jul/2026", "0.2",
     "CR-001 (D-CR1 to D-CR7): domain model and data view updated for the two-value source, the "
     "write-once severity rule, the category-weights plus trust-slider profile, risk as a scoring "
     "multiplier, and the in-place finding detail.", "Group 16"),
    ("31/Jul/2026", "0.3",
     "CR-001 (D-CR8 to D-CR12): the data view stores facts only and every score is derived on read; "
     "the category taxonomy is fixed at six values, so the profile carries six weights.", "Group 16"),
    ("01/Aug/2026", "0.4",
     "New Section 6.2 - the apply-profile write path, contrasted with the scan path. The active "
     "profile becomes workspace state so the read endpoints stay unparameterised.", "Group 16"),
    ("03/Aug/2026", "0.5",
     "Aligned with SRS v1.0. Three decisions recorded: repository access in v1.0 is public URL paste "
     "only, with the GitHub App deferred to v1.1; scoring is a pure Python function and is never "
     "computed in the database; there are no webhooks and no role-based access control in v1.0. "
     "Figures numbered 1 to 10.", "Group 16"),
    ("05/Aug/2026", "0.6",
     "Section 6.2 gains the step-by-step apply-profile workflow and the reason for PUT over PATCH. "
     "Section 9 adds the rule separating a cache from a source of truth, with the cache key defined "
     "on the inputs rather than the session. Section 10 adds the read-cost arithmetic, the exact "
     "per-group aggregate that makes the trend chart cheap, and the four-stage caching ladder with "
     "the ETag at stage 1. No architectural decision changed.", "Group 16"),
]
import copy as _copy
while len(rev.rows) < 1 + len(rows):
    rev._tbl.append(_copy.deepcopy(rev.rows[-1]._tr))
for ri, data in enumerate(rows, start=1):
    for ci, val in enumerate(data):
        cell = rev.rows[ri].cells[ci]
        para = cell.paragraphs[0]
        para.style = doc.styles["Tabletext"]
        if para.runs:
            para.runs[0].text = val
            for extra in para.runs[1:]:
                extra.text = ""
        elif val:
            para.add_run(val)

# ───────────────────────── 4. drop the figure-guidance block ─────────────────
# Everything between the revision table and the TOC field is RUP guidance.
kids = list(body.iterchildren())
rev_idx = next(i for i, e in enumerate(kids) if e.tag == qn('w:tbl'))
toc_start = next(i for i, e in enumerate(kids)
                 if e.tag == qn('w:p') and 'TOC \\o' in e.xml)
for e in kids[rev_idx + 1:toc_start]:
    body.remove(e)

# ───────────────────────── 5. TOC: fresh auto-updating field ─────────────────
kids = list(body.iterchildren())
toc_start = next(i for i, e in enumerate(kids)
                 if e.tag == qn('w:p') and 'TOC \\o' in e.xml)
intro = next(i for i, e in enumerate(kids)
             if e.tag == qn('w:p') and 'Heading1' in e.xml and 'Introduction' in e.xml)
after = kids[toc_start - 1]
for e in kids[toc_start:intro]:
    body.remove(e)

toc_xml = (
    '<w:p %s><w:pPr><w:pStyle w:val="TOC1"/></w:pPr>'
    '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
    '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    '<w:r><w:t xml:space="preserve">Right-click here and choose "Update Field" '
    '(or press Ctrl+A then F9) to generate the Table of Contents.</w:t></w:r>'
    '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>' % nsdecls('w')
)
after.addnext(parse_xml(toc_xml))

# ───────────────────────── 6. wipe the old body content ──────────────────────
kids = list(body.iterchildren())
start = next(i for i, e in enumerate(kids)
             if e.tag == qn('w:p') and 'Introduction' in e.xml and 'Heading1' in e.xml)
for e in kids[start:]:
    if e.tag != qn('w:sectPr'):
        body.remove(e)

# ───────────────────────── 7. content helpers ────────────────────────────────
def h1(t):
    doc.add_paragraph(t, style="Heading 1")


def h2(t):
    doc.add_paragraph(t, style="Heading 2")


def h3(t):
    doc.add_paragraph(t, style="Heading 3")


def info(t):
    doc.add_paragraph(t, style="InfoBlue")


def body_text(t, bold_lead=None):
    p = doc.add_paragraph(style="Body Text")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(t)
    return p


def bullet(t, lead=None):
    p = doc.add_paragraph(style="Bullet1")
    p.add_run("\u2022\t")
    if lead:
        r = p.add_run(lead)
        r.bold = True
    p.add_run(t)
    return p


def caption(t):
    p = doc.add_paragraph(style="Body Text")
    r = p.add_run(t)
    r.bold = True
    r.italic = True
    return p


def _fill(cell, text, bold=False, center=False):
    cell._tc.remove(cell._tc.find(qn('w:p')))
    for line in str(text).split("\n"):
        p = cell.add_paragraph(style="Tabletext")
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j, seg in enumerate(re.split(r"\*\*(.+?)\*\*", line)):
            if seg:
                r = p.add_run(seg)
                r.bold = bold or (j % 2 == 1)


def table(headers, data, widths, cap=None):
    """Table with a filled header row and `data` rows (may contain '')."""
    if cap:
        caption(cap)
    t = doc.add_table(rows=1 + len(data), cols=len(headers))
    tbl = t._tbl
    tbl.replace(tbl.tblPr, parse_xml(TBLPR))
    grid = tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(w))
    for ci, htxt in enumerate(headers):
        c = t.rows[0].cells[ci]
        c.width = docx.shared.Twips(widths[ci])
        _fill(c, htxt, bold=True, center=True)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml('<w:tblHeader %s/>' % nsdecls('w')))
    for ri, rowdata in enumerate(data, start=1):
        for ci, val in enumerate(rowdata):
            c = t.rows[ri].cells[ci]
            c.width = docx.shared.Twips(widths[ci])
            _fill(c, val)
    doc.add_paragraph(style="Body Text")
    return t


def grid(headers, widths, blank_rows=4, cap=None, first_col=None):
    """Empty grid to be filled in by the team.

    first_col: optional list of pre-filled left-column labels (structure, not
    content) - the remaining columns are left blank.
    """
    if first_col:
        data = [[lbl] + [""] * (len(headers) - 1) for lbl in first_col]
    else:
        data = [[""] * len(headers) for _ in range(blank_rows)]
    return table(headers, data, widths, cap)


FIGNOTE = ("[Insert Figure %d here, then describe it below in two or three sentences that "
           "begin with \"Figure %d ...\". Black text 12pt or larger, white background, "
           "in line with text.]")


def figure(num, title, extra=""):
    info((FIGNOTE % (num, num)) + (" " + extra if extra else ""))
    caption("Figure %d. %s" % (num, title))
    doc.add_paragraph(style="Body Text")


W2 = [2300, 7060]
W2E = [4680, 4680]
W3 = [2400, 3480, 3480]
W3B = [1200, 4080, 4080]
W4 = [1000, 2600, 2600, 3160]

# ═════════════════════════ 1. INTRODUCTION (written) ═════════════════════════
h1("Introduction")

h2("Purpose")
body_text(
    "This document describes the software architecture of Code Sage AI, the AI-powered "
    "technical-debt analytics dashboard specified in the Software Requirements Specification. "
    "It uses the RUP \"4+1\" model, so the system is presented through five views - Use-Case, "
    "Logical, Process, Deployment and Implementation - together with a Data View, because the "
    "persistent schema in this system carries real architectural weight and leaving it out would "
    "hide a central decision.")
body_text(
    "The document has three audiences. The development team uses it to know which component owns "
    "which responsibility before writing code. The mentor and the evaluators use it to judge "
    "whether the design actually delivers what the SRS promises. Future maintainers use it to "
    "understand why the system is shaped this way, so that a later change does not quietly break "
    "an assumption the design depends on.")
body_text(
    "It should be read together with the SRS, which states what the system must do. This document "
    "states how the parts are arranged to do it. Where the two overlap, the SRS is the normative "
    "source, and this document refers back to the requirement by its identifier.")

h2("Scope")
body_text(
    "This document covers the architecture of the v1.0 release: a multi-tenant web application in "
    "which a user signs in with GitHub, connects a public repository by pasting its URL, runs a "
    "scan on a chosen branch when they decide to, and reads a prioritised technical-debt dashboard "
    "built from the stored result. It covers all four parts of the system - the frontend, the "
    "backend API and its asynchronous workers, the machine-learning service, and the database.")
body_text("Three limits are stated here because they shape several of the views that follow.")
bullet(
    "Sign-in uses GitHub OAuth, and the analysis pipeline reads the repository through an ordinary "
    "git clone. There is no GitHub App installation, no private-repository support and no webhook "
    "endpoint in v1.0. Those belong to later releases, and this document marks the points where "
    "they would attach (SRS FR-3, Table 3-21).",
    lead="Repository access in v1.0 is public repositories by URL only. ")
bullet(
    "There is no automatic or event-driven analysis in v1.0. A scan happens because a user asked "
    "for it (SRS FR-6).",
    lead="Scans are started by the user, never by an event. ")
bullet(
    "Team management, role-based access control, private repositories and additional Git hosts are "
    "mentioned only where they affect a v1.0 decision - for example the tenant column that exists "
    "from day one so that multi-user workspaces can be added later without a migration (SRS FR-2, "
    "FR-23).",
    lead="Later-release concerns are out of scope. ")

h2("Definitions, Acronyms, and Abbreviations")
body_text(
    "The project keeps a single glossary in SRS Section 1.3, and that glossary is the authority. "
    "The terms used most often in this document are repeated below for the reader's convenience.")
table(["Term", "Meaning in this document"],
      [("Finding",
        "One detected issue at a file:line:symbol, carrying a source, a category, a severity and a "
        "one-line reason. The atomic unit of output."),
       ("source",
        "Which detector produced a finding. Exactly two values: rule or satd (SRS FR-8.2)."),
       ("category",
        "What kind of debt a finding is. Exactly six values: code-design, requirement, defect, "
        "documentation, test and security (SRS FR-9.3)."),
       ("Snapshot",
        "The stored, immutable result of one scan, identified by repository, branch, commit SHA "
        "and time."),
       ("Scoring profile",
        "The six category weights plus the trust slider that turn stored findings into scores "
        "(SRS FR-20)."),
       ("Visibility floor",
        "The rule that a critical security finding stays visible no matter how the profile is set "
        "(SRS FR-24)."),
       ("Workspace / tenant",
        "The top-level owner of data. A project belongs to a workspace, not to a person."),
       ("RLS",
        "PostgreSQL Row-Level Security, the mechanism used to keep tenants apart."),
       ("ML-1 / ML-2",
        "The SATD comment classifier and the file-level bug-proneness risk model.")],
      W2, "Table 1-1. Terms used most often in this document.")

h2("References")
body_text("The documents referenced elsewhere in this SAD are listed below.")
table(["#", "Document", "Where it is held"],
      [("1", "Software Requirements Specification, Code Sage AI v1.0. The normative source for "
             "every requirement identifier used in this document.",
        "docs/Deliverables/Software_Requirements_Specification.docx"),
       ("2", "Change Request CR-001 - scoring model and finding UX (decisions D-CR1 to D-CR12).",
        "docs/Change Requests/"),
       ("3", "Backend Analysis Engine design note.",
        "docs/Project Management & Planning/"),
       ("4", "Release Roadmap and Frontend Prototype Plan.",
        "docs/Project Management & Planning/"),
       ("5", "Project Proposal and Feasibility Report.",
        "docs/Deliverables/"),
       ("6", "Shared data contract - the single source of truth for the shapes crossing the "
             "frontend, backend and database boundary (SRS SP-4).",
        "apps/web/src/lib/types/index.ts")],
      W3B, "Table 1-2. Referenced documents.")
info("[TO FILL - name the tool used to draw the ten figures (for example draw.io) and add it here "
     "as an IEEE-style reference with an accessed date, as required by the template and by SRS "
     "Appendix D. Tool and library references such as Lizard, PyDriller, scikit-learn, FastAPI, "
     "Redis, Celery, PostgreSQL, Next.js and Tailwind CSS are already listed in SRS Section 1.4 "
     "and do not need repeating here.]")

h2("Overview")
body_text("The rest of the document is organised as follows.")
body_text(
    "Section 2 explains which architectural views are used and why. Section 3 lists the goals and "
    "constraints that actually shaped the design, each with its architectural consequence. Sections "
    "4 to 8 present the five views: the use cases that exercise the architecture, the decomposition "
    "into subsystems and classes, the runtime behaviour of the processes, the physical deployment, "
    "and the layering of the codebase. Section 9 describes the persistent data model and the rule "
    "that separates stored facts from derived scores. Section 10 covers sizing and performance "
    "targets, Section 11 explains how the architecture delivers the quality attributes, and Section "
    "12 lists the references.")
body_text(
    "Ten figures are referenced by number throughout. Each one is cited in the text as \"Figure n\" "
    "and is described in the sentences that follow it.")

# ═════════════════════════ 2. ARCHITECTURAL REPRESENTATION ═══════════════════
h1("Architectural Representation")
info("[TO FILL - say in two or three sentences what kind of system this is (a multi-tenant web "
     "application with an asynchronous analysis pipeline and an offline-trained machine-learning "
     "component), then complete the table below. The Data View is normally optional; state why it "
     "is included here - the schema encodes the rule about which values are stored and which are "
     "computed. Close the section with the one invariant that runs through every view: detection "
     "and scoring happen on the server, and the dashboard only reads and draws.]")
grid(["View", "What it captures", "Diagrams (figure numbers)"], W3,
     cap="Table 2-1. Architectural views used in this document.",
     first_col=["Use-Case", "Logical", "Process", "Deployment", "Implementation", "Data"])

# ═════════════════════════ 3. GOALS AND CONSTRAINTS ══════════════════════════
h1("Architectural Goals and Constraints")
info("[TO FILL - list only the goals that actually changed a design decision, and for each one say "
     "what it forced in the architecture. Cover at least: secure multi-tenancy, keeping the UI "
     "responsive during a scan, low-noise and explainable output, changing a profile without a "
     "re-scan, severity never being influenced by the user, the prototype becoming the product, "
     "core-first delivery, and the containerised free stack. Cite the driving requirement in the "
     "right-hand column.]")
grid(["#", "Goal or constraint", "What it forces in the architecture", "Driving requirement"], W4,
     cap="Table 3-1. Architectural goals and their consequences.",
     first_col=["G1", "G2", "G3", "G4", "G5", "G6", "G7", "G8"])
info("[TO FILL - follow the table with a short paragraph on the fixed constraints: the technology "
     "stack, supervised learning only, models trained offline on public datasets, a three-member "
     "team and an academic timeline. Explain in one sentence why v1.0 is one complete vertical "
     "slice rather than several partial features.]")

# ═════════════════════════ 4. USE-CASE VIEW ══════════════════════════════════
h1("Use-Case View")
info("[TO FILL - name the actors first: Developer (primary), Tech Lead or Manager, Viewer or "
     "Stakeholder, Org Admin (v2 only), GitHub as an external system, and the ML Service as an "
     "internal system actor. Then complete the table. The template requires at least five use "
     "cases; eight are listed below. In the last column say why each one is architecturally "
     "significant, not what it does.]")
grid(["ID", "Use case", "Actors", "Why it is architecturally significant"], W4,
     cap="Table 4-1. Architecturally significant use cases.",
     first_col=["UC-1", "UC-2", "UC-3", "UC-4", "UC-5", "UC-6", "UC-7", "UC-8"])
figure(1, "Use-case overview for v1.0.",
       "Draw it as a proper UML use-case diagram: ovals inside a labelled system boundary, with "
       "the actors outside it.")

h2("Use-Case Realizations")
info("[TO FILL - complete one realization table for each significant use case. Three blank tables "
     "are provided below, for UC-3 Run a scan, UC-2 Connect a repository and UC-6 Select a scoring "
     "profile. UC-3 is the central one because it exercises every layer, every process and every "
     "external interface, so give it the most detail. Write the main flow as numbered steps that "
     "name the internal objects involved, never just \"user\" and \"system\". Use N/A where an item "
     "does not apply.]")

REAL_ROWS = ["Use case name", "Actors", "Description", "Preconditions", "Main flow",
             "Successful end / post condition", "Fail end / post condition", "Extensions"]
for n, who in ((2, "UC-3 Run a scan"), (3, "UC-2 Connect a repository"),
               (4, "UC-6 Select a scoring profile")):
    grid(["Item", "Detail"], W2E,
         cap="Table 4-%d. Use-case realization: %s." % (n, who),
         first_col=REAL_ROWS)

# ═════════════════════════ 5. LOGICAL VIEW ═══════════════════════════════════
h1("Logical View")

h2("Overview")
info("[TO FILL - describe the decomposition into the four subsystems and the layers inside each "
     "one, then insert Figure 2. In the description, point out where the scoring engine sits: it "
     "is called by the domain services on the read path, not by the workers on the write path.]")
grid(["Subsystem", "Responsibility", "Main internal parts"], W3,
     cap="Table 5-1. Subsystem decomposition.",
     first_col=["Frontend (Next.js)", "Backend (FastAPI and Celery)",
                "ML Service (Python, scikit-learn)", "Database (PostgreSQL)"])
figure(2, "Logical decomposition into subsystems and layers.")

h2("Architecturally Significant Design Packages")
info("[TO FILL - insert the class diagram as Figure 3, then complete the table for the classes that "
     "matter architecturally. After the table, write out the three rules the rest of the design "
     "depends on: (a) category and severity are written once by whoever detected the finding and "
     "are never recomputed (SRS FR-8.1, FR-9.2); (b) source has exactly two values, rule and satd, "
     "because only two components emit findings and security is a category rather than a source "
     "(SRS FR-8.2); (c) the scoring profile holds six category weights and one trust value, and "
     "never holds severities (SRS FR-20).]")
figure(3, "Core domain model.")
grid(["Class", "Responsibility", "Key attributes", "Notes"], W4,
     cap="Table 5-2. Architecturally significant classes.",
     first_col=["Workspace", "Repo", "Branch", "Scan", "Finding", "FileScore", "ScoreProfile",
                "ScoringEngine"])

# ═════════════════════════ 6. PROCESS VIEW ═══════════════════════════════════
h1("Process View")
info("[TO FILL - name the four kinds of process in a running system and say how they communicate: "
     "the stateless API process, one or more Celery worker processes, the ML inference process, and "
     "the database. The API and the workers talk through the Redis broker; everything persists to "
     "PostgreSQL. Complete the table, then insert the activity diagram as Figure 4 and point out "
     "that the write path ends at \"persist the snapshot\" - scoring is not a pipeline stage, it "
     "happens later on the read path.]")
grid(["Process", "Kind", "Talks to", "What it carries"], W4,
     cap="Table 6-1. Processes and their communication.",
     first_col=["API (FastAPI)", "Worker (Celery)", "ML inference", "PostgreSQL"])
figure(4, "Scan activity, including skip-if-unchanged and cancel.")

h2("The extraction boundary")
info("[TO FILL - state the boundary in one line: git history enters the pipeline as numbers, never "
     "as text (SRS FR-7.1). Explain why - text produces findings, and a finding must land on a "
     "file:line the user can open, so text is read from the checked-out tree; history produces "
     "metrics, so it may look backwards, but only as a numeric feature vector. Complete the table, "
     "then write the two consequences: a scan is a pure function of the repository at one commit, "
     "which is what makes skip-if-unchanged safe; and the 90-day churn window is measured backwards "
     "from the scanned commit's committer date, never from the clock, or the same commit would "
     "score differently on different days (Decision D6; SRS FR-11).]")
grid(["Input", "Used at scan time?", "Why"], W3,
     cap="Table 6-2. What may and may not enter the pipeline.",
     first_col=["Source comments at the scanned SHA", "Source files at the scanned SHA",
                "Commit history reachable from that SHA", "Commit message text",
                "Pull requests, issues, GitHub API metadata", "Previously stored snapshots"])
figure(5, "Run-a-scan sequence.",
       "Show the internal objects, not a black box: DashboardUI, ScanControl, the FastAPI process, "
       "the Redis broker, the Celery worker, the extractors, the rule engine, the ML service, the "
       "scoring engine and the database. Make clear that the scoring engine is invoked by the API "
       "on the read path, never by the worker.")

h2("Applying a scoring profile")
info("[TO FILL - explain that a profile change is the only other user-initiated write in v1.0 and "
     "is deliberately the opposite of a scan: nothing is queued, no worker wakes and no snapshot is "
     "written (SRS FR-20, FR-21). Then complete the workflow table below step by step.]")
grid(["Step", "What happens"], W2E,
     cap="Table 6-3. Applying a profile, step by step.",
     first_col=[
         "1. User drags the sliders",
         "2. User clicks Apply",
         "3. PUT /api/profiles/active",
         "4. Server validates and clamps",
         "5. Server writes",
         "6. Server responds",
         "7. Client re-reads the dashboard",
         "8. Server scores and returns"])
info("[TO FILL - after the table, answer two questions a reader will have. First, why PUT and not "
     "PATCH: PUT means make this resource look exactly like this, so the body is the whole profile "
     "rather than a change to it, and a retry after a dropped response cannot double-apply. That "
     "matters because step 7 fires immediately after step 6, and a half-applied profile would "
     "render a dashboard matching no profile that exists. Second, why no profile appears in the URL "
     "of step 7: the active profile is workspace state on the server, so if it travelled as a query "
     "parameter every read endpoint would grow a parameter each time the profile shape changed, and "
     "the scoring formula would leak into the API surface.]")
grid(["Aspect", "Run a scan", "Apply a profile"], W3,
     cap="Table 6-4. The two write paths compared.",
     first_col=["Request", "Handled by", "Writes", "Duration", "Effect on findings"])
figure(6, "Apply-a-profile sequence.")
info("[TO FILL - close the section with the three properties this shape buys - the read endpoints "
     "stay unparameterised, the PUT is idempotent by construction, and the profile is shared rather "
     "than per-tab so the name on the trend chart always refers to something real - and with why "
     "clamping is a server-side rule and not a client convenience.]")

h2("The visibility floor")
info("[TO FILL - SRS FR-24 requires a critical security finding to stay visible however the profile "
     "is set. Complete the table below with the three mechanisms that deliver it, and say which are "
     "structural (so they cannot be forgotten) and which is an explicit step in the scoring engine "
     "that needs its own test case.]")
grid(["#", "Mechanism", "Why it cannot be bypassed"], W3B,
     cap="Table 6-5. The three mechanisms of the visibility floor.",
     first_col=["1", "2", "3"])

# ═════════════════════════ 7. DEPLOYMENT VIEW ════════════════════════════════
h1("Deployment View")
info("[TO FILL - insert the deployment diagram as Figure 7 and complete the node table. Say which "
     "containers scale and why: the workers are the only ones whose load grows with usage, because "
     "concurrent scans are the growing load. Note that each worker needs local disk for the clone "
     "it is analysing, about 2 GB per concurrent scan, released when the scan ends (SRS Table "
     "3-20). Mark Redis and PostgreSQL as private-network only.]")
figure(7, "Deployment for v1.0.")
grid(["Node", "Processes hosted", "Connections"], W3,
     cap="Table 7-1. Physical nodes and the processes mapped onto them.",
     first_col=["Client device", "Frontend container", "Backend container", "Worker container(s)",
                "ML container", "Redis broker", "PostgreSQL"])
info("[TO FILL - confirm the hosting provider (the Feasibility Report cites DigitalOcean) and add "
     "the TLS termination and domain nodes to Figure 7.]")

# ═════════════════════════ 8. IMPLEMENTATION VIEW ════════════════════════════
h1("Implementation View")

h2("Overview")
info("[TO FILL - describe the monorepo: three applications (apps/web, apps/api, apps/ml) organised "
     "into three layers, presentation then application/domain then data, with the shared data "
     "contract as the only agreed shape crossing the frontend and backend boundary. State the "
     "layering rule plainly - the presentation layer never reaches the database directly - then "
     "insert Figure 8 and point out that the dependency arrows only ever point downward.]")
figure(8, "Implementation layers and their components.")
grid(["Layer", "What belongs in it", "Rule for inclusion"], W3,
     cap="Table 8-1. The three layers.",
     first_col=["Presentation", "Application / Domain", "Data"])

h2("Layers")
info("[TO FILL - one short paragraph per layer naming its subsystems. For the application layer, "
     "state explicitly that ScoringEngine is a pure Python function taking stored findings, "
     "per-file facts and the active profile, and returning priorities, file debt, health, grade and "
     "the category breakdown - it is not computed in the database. Then complete the table of layer "
     "rules and insert the package diagram as Figure 9.]")
grid(["Rule", "What it requires", "Requirement"], W3,
     cap="Table 8-2. Rules that govern the layers.",
     first_col=["Volatile libraries sit behind one boundary each",
                "Thresholds and weights are configuration, not code literals",
                "Models are versioned artifacts loaded at runtime",
                "The presentation layer never reaches the database"])
figure(9, "Package diagram of the backend application.",
       "Show the internal packages of apps/api - routers, services, scoring, tasks, extractors and "
       "db - and the dependencies between them.")

# ═════════════════════════ 9. DATA VIEW ══════════════════════════════════════
h1("Data View")
info("[TO FILL - insert the entity-relationship diagram as Figure 10 and complete the entity table. "
     "Note that RLS policies key every tenant-owned table on workspace_id, and that SCAN rows are "
     "append-only so trend, history and delta are queries over existing rows rather than updates to "
     "them.]")
figure(10, "Multi-tenant data model.")
grid(["Entity", "What it holds", "Stored fact or derived?"], W3,
     cap="Table 9-1. Persistent entities.",
     first_col=["WORKSPACE", "USER / MEMBERSHIP", "REPO", "BRANCH", "SCAN", "FINDING",
                "FILE_SCORE", "SCORE_PROFILE"])
info("[TO FILL - this is the most important rule in the data view, so write it out in full (SRS "
     "FR-21, SP-7). The schema stores facts, not scores. FINDING keeps the evidence; FILE_SCORE "
     "keeps the two per-file inputs, risk_score from ML-2 and churn_factor from the commit-anchored "
     "window. All of these are properties of the code at that commit. Priority, file debt, health "
     "score, grade, delta and the category breakdown are functions of the active profile, so they "
     "are computed on every read and are never columns. Explain what would go wrong without this "
     "rule: an editable profile would leave every stored score stale the moment a slider moved.]")
info("[TO FILL - record the decision that scoring is computed in Python and not in SQL, with the "
     "three reasons, so that it is not quietly reversed later. First, the formula in SRS FR-11 has "
     "five factors, bounded ranges and a visibility-floor override, and in SQL it becomes hard to "
     "read and harder to change. Second, SRS SP-11 requires the scoring path to be deterministic "
     "and exactly testable, and a Python function can be unit-tested against the worked example in "
     "SRS TC-11 with no database at all. Third, SRS SP-8 requires thresholds and weights to live in "
     "configuration, which a Python function reads naturally but a SQL view would turn into a "
     "migration.]")
grid(["Column or table", "Rule that applies to it", "Requirement"], W3,
     cap="Table 9-2. Columns and tables with a rule attached.",
     first_col=["FINDING.severity and FINDING.category", "FINDING.source",
                "SCORE_PROFILE.weights and trust_s", "WORKSPACE.active_profile_id",
                "SCAN (append-only)"])
info("[TO FILL - deriving every score on read raises the question of whether the same answer is "
     "being recomputed over and over. Caching it is allowed; what matters is what kind of thing is "
     "stored. Write out the rule and the cache key. Because scoring is a pure function over "
     "immutable snapshot facts, the same inputs always produce the same output, so a cache needs no "
     "invalidation logic - a stale answer is impossible. The key is (scan_id, profile_fingerprint), "
     "where the fingerprint is a hash of the six weights and trust_s. State explicitly that the key "
     "is the inputs and never the session: a session-keyed cache would put derived state in server "
     "memory, which SRS FR-1 and FR-21 both forbid, and two tabs would compute the same result "
     "twice. Then give the rule that keeps this honest - a cached value may be deleted at any "
     "moment without losing information, because it can be rebuilt from stored facts; the moment a "
     "derived value would be missed if deleted, it has become a second source of truth, and that is "
     "what FR-21 forbids. Complete the table below and confirm that priority, debt score, health "
     "score, grade, delta and the category breakdown are still never columns.]")
grid(["Value", "Cache or source of truth?", "Why"], W3,
     cap="Table 9-3. What may be cached and what may not.",
     first_col=["Computed HealthReport (ETag or Redis entry)",
                "Per-group aggregate sums (two per category-source group per scan)",
                "priority, debt_score, health_score, grade, delta, category breakdown",
                "FINDING evidence, FILE_SCORE risk and churn"])
info("[TO FILL - list what is deliberately NOT in the v1.0 schema and why, so a reader does not "
     "think it was forgotten: no SUPPRESSION table and no finding-action history, because v1.0 is "
     "view-only (SRS FR-17c); FINDING.status exists and every v1.0 finding is open, which is what "
     "SRS FR-11 means by the sum of open finding priorities; no webhook-event table, because scans "
     "are manual; and no role or permission tables beyond MEMBERSHIP.role, because RBAC is v2. Then "
     "finalise the column types and the indexes - at minimum on (repo_id, branch, scanned_at) for "
     "the trend query and on (scan_id, category) for the breakdown - and write out the exact RLS "
     "policies.]")

# ═════════════════════════ 10. SIZE AND PERFORMANCE ══════════════════════════
h1("Size and Performance")
info("[TO FILL - state the baseline the system is sized for, then explain in four short points how "
     "the architecture meets it: analysis never blocks the API because scans run on workers; reads "
     "are cheap because a dashboard read is one query plus an in-memory scoring pass; workers scale "
     "horizontally; and GitHub rate limits are avoided rather than managed, because the pipeline "
     "reads through git clone and uses no REST quota. Then complete the table.]")
grid(["Characteristic", "Target", "Source requirement"], W3,
     cap="Table 10-1. Dimensioning characteristics and performance targets.",
     first_col=["Registered users at baseline", "Interaction feedback", "Non-analysis interactions",
                "Scan enqueue", "Progress reporting", "Scan progress polling", "Worker disk"])
h2("What a dashboard read costs")
info("[TO FILL - answer the obvious objection to deriving scores on read, using the operation count "
     "rather than an opinion. Per finding the formula is four multiplications and three lookups; "
     "the findings are then grouped and summed by file, and repo health is one division, so the "
     "work is linear in the number of findings in one snapshot. Complete the table with your own "
     "measured figures. Then state the consequence that matters: reading and deserialising the rows "
     "out of PostgreSQL costs more than the arithmetic performed on them, so the database read "
     "dominates and optimising the multiplication would be optimising the wrong half. Compare the "
     "result against the one-second budget in SRS PERF-02.]")
grid(["Repository size", "Findings, realistically", "Scoring time in Python"], W3,
     cap="Table 10-2. Cost of scoring one snapshot.",
     first_col=["Typical target repository (about 50 KLOC)", "Large repository",
                "Pessimistic upper bound"])

h2("The trend chart")
info("[TO FILL - this is the one place where the cost is real, so explain it and give the fix. SRS "
     "FR-14 redraws every point of the trend under the active profile, so fifty snapshots of a "
     "thousand findings means fifty thousand findings read for one chart. Show why the fix is exact "
     "rather than approximate: within one category-and-source group the category weight and the "
     "source trust are constant, and the two inner sums - base times churn, and base times churn "
     "times risk - contain no profile value at all, only stored severity, churn and risk. So two "
     "sums per group per scan reproduce the total exactly under any profile. With six categories "
     "and two sources that is twelve groups, so twenty-four numbers per snapshot, and a fifty-point "
     "trend becomes about twelve hundred multiply-adds instead of fifty thousand row reads. Refer "
     "back to Section 9: these sums are facts rather than scores, so storing them does not conflict "
     "with SRS FR-21.]")

h2("Caching, cheapest measure first")
info("[TO FILL - complete the table below, then explain the two stages worth describing in detail. "
     "Stage 1 is worth building in v1.0 because it costs almost nothing: the server computes an "
     "ETag from a hash of the scan identifier and the profile fingerprint using two small lookups, "
     "without reading a single finding, and a browser sends If-None-Match automatically on refresh, "
     "so when neither the snapshot nor the profile has changed the server replies 304 Not Modified "
     "with an empty body - no findings read, no scoring performed, no payload sent, and no "
     "server-side state added. Stage 2 reuses the Redis instance already present for Celery, so a "
     "hit means a result another user or tab already computed is reused rather than recalculated. "
     "Close by confirming that all four stages remain consistent with SRS FR-21 and SP-7, because "
     "every one of them stores something rebuildable from the snapshot facts.]")
grid(["Stage", "Measure", "When to build it"], W3,
     cap="Table 10-3. Caching measures, in the order they should be considered.",
     first_col=["0", "1", "2", "3"])

info("[TO FILL - four numbers must be measured, not guessed, before the final submission: scan time "
     "per KLOC, maximum concurrent scans per worker, p95 dashboard read latency, and database size "
     "per snapshot. Record the measured values and the conditions they were measured under.]")

# ═════════════════════════ 11. QUALITY ═══════════════════════════════════════
h1("Quality")
info("[TO FILL - for each attribute, say what the architecture actually does to deliver it, not "
     "that it is important. Cite the requirement each claim answers to. Testability and "
     "traceability matter here as much as security, because they are what make the deterministic "
     "detection path and the pure scoring function verifiable.]")
grid(["Attribute", "How the architecture delivers it", "Requirement"], W3,
     cap="Table 11-1. Quality attributes and their architectural support.",
     first_col=["Security and privacy", "Reliability", "Scalability",
                "Maintainability and extensibility", "Usability", "Portability", "Testability",
                "Traceability and diagnosability"])

# ═════════════════════════ 12. REFERENCES ════════════════════════════════════
h1("References")
info("[TO FILL - list the references in IEEE style. The document-level references are already in "
     "Table 1-2. Cite tools by their official web page with an accessed date. Cite algorithms, "
     "techniques and theories from textbooks or peer-reviewed articles, and similar work from "
     "research articles. The diagram tool must be named here.]")
doc.add_paragraph(style="Body Text")

doc.save(OUT)

# ───────────────────────── 8. make Word refresh the TOC on open ──────────────
tmp = OUT + ".tmp"
with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if re.match(r"word/(header|footer)\d+\.xml$", item.filename):
            # placeholders inside w:fldSimple are not reachable through python-docx runs
            s = data.decode("utf-8")
            for k, v in REPL.items():
                s = s.replace(k.replace("<", "&lt;").replace(">", "&gt;"), v)
            data = s.encode("utf-8")
        if item.filename == "word/settings.xml":
            s = data.decode("utf-8")
            if "updateFields" not in s:
                for anchor in ("<w:hdrShapeDefaults", "<w:footnotePr", "<w:endnotePr",
                               "<w:compat", "<w:rsids", "</w:settings>"):
                    if anchor in s:
                        s = s.replace(anchor, '<w:updateFields w:val="true"/>' + anchor, 1)
                        break
            data = s.encode("utf-8")
        zout.writestr(item, data)
os.replace(tmp, OUT)

d2 = docx.Document(OUT)
print("OK ->", OUT)
print("paragraphs:", len(d2.paragraphs), " tables:", len(d2.tables), " sections:", len(d2.sections))
print("size KB:", round(os.path.getsize(OUT) / 1024, 1))
